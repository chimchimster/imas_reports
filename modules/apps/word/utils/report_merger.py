import os
import re
import docx

from docxcompose.composer import Composer

from modules.logs.decorators import tricky_loggy
from modules.folder import FolderUUID


class MergeReport:
    """ Класс обращается к заданным директориям
        основываясь на уникальных идекнтификаторах клиента uuid
        и собирает воедино сгенерированные части шаблонов. """

    storage_api = os.getenv('storage_api')

    def __init__(self):
        self.path_to_folder: str = os.path.join(
            os.getcwd(),
            'word',
            'temp',
        )
        self.path_to_templates: str = os.path.join(
            os.getcwd(),
            'word',
            'temp_templates',
        )
        self.path_to_result: str = os.path.join(
            os.getcwd(),
            'word',
            'merged',
        )
        self.folder: FolderUUID = None

    @tricky_loggy
    def set_path_to_folder(self, folder: FolderUUID):
        self.path_to_folder += f'/{folder.unique_identifier}'

    @tricky_loggy
    def set_path_to_result(self, folder: FolderUUID):
        self.path_to_result += f'/{folder.unique_identifier}'

    @tricky_loggy
    def set_path_to_templates(self, folder):
        self.path_to_templates += f'/{folder.unique_identifier}'

    @tricky_loggy
    def create_result_folder(self):
        os.chdir(
            os.path.join(
                os.getcwd(),
                'word',
                'merged',
            ),
        )

        if not os.path.exists(
            os.path.join(
                os.getcwd(),
                'word',
                'merged',
                f'{self.folder.unique_identifier}',
            )
        ):
            os.mkdir(f'{self.folder.unique_identifier}')

        os.chdir('../../../..')

    @tricky_loggy
    def merge(self) -> None:

        self.set_path_to_folder(self.folder)
        self.set_path_to_templates(self.folder)

        master: docx.Document = docx.Document(
            os.path.join(
                self.path_to_templates,
                'out.docx',
            )
        )
        composer: Composer = Composer(master)

        file_order = [file for file in os.listdir(self.path_to_folder)]
        file_order.sort(key=lambda file_name: int(re.findall(r'[0-9]+', file_name)[0]))

        for idx, file in enumerate(file_order):

            file_path = os.path.join(
                self.path_to_folder,
                file,
            )

            if os.path.isfile(file_path) and file.endswith('.docx'):
                doc = docx.Document(file_path)

                if file_path.endswith('table.docx'):
                    run = master.add_paragraph().add_run()
                    run.add_break(docx.enum.text.WD_BREAK.PAGE)

                if file_path.endswith('Process.docx'):

                    last_paragraph = master.paragraphs[-1]
                    last_run = last_paragraph.runs[-1]
                    last_run.add_break(docx.enum.text.WD_BREAK.PAGE)

                if idx == 1:
                    run = master.add_paragraph().add_run()
                    run.add_break(docx.enum.text.WD_BREAK.PAGE)

                composer.append(doc)

        self.create_result_folder()
        self.set_path_to_result(self.folder)

        output_file = os.path.join(
            self.path_to_result,
            f'{self.folder.unique_identifier}.docx',
        )

        composer.save(output_file)
