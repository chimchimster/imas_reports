import os
import uuid
import docx
import shutil

from docxtpl import DocxTemplate, InlineImage
from docx.shared import RGBColor, Pt, Cm
from docxcompose.composer import Composer

from multiprocessing import Semaphore, Process

from modules.logs.decorators import tricky_loggy
from modules.apps.localization import ReportLanguagePicker
from modules.apps.word.utils.tools.auxiliary_functions import generate_chart_categories
from modules.mixins import AbstractRunnerMixin
from modules.apps.word.tools import TableStylesGenerator, SchedulerStylesGenerator, HighchartsCreator, MetricsGenerator
from modules.apps.word.utils.tools.decorators import render_diagram, throw_params_for_distribution_diagram, render_map


class TableProcess(AbstractRunnerMixin):

    def __init__(self, proc_object, data, report_format):
        super().__init__(proc_object, data, report_format)
        self._template_path: str = os.path.join(
            os.getcwd(),
            'word',
            'temp_templates',
            f'{self.proc_obj.folder.unique_identifier}',
            'template_parts',
            'table.docx',
        )

    @tricky_loggy
    def create_temp_template_folder(self) -> None:

        _type_of_table = self.proc_obj.type

        _uuid: str = '_'.join((_type_of_table, str(self.proc_obj.folder.unique_identifier)))

        path_to_temp_table_folder = os.path.join(
            os.getcwd(),
            'word',
            'temp_tables',
            'templates',
            _uuid,
        )

        if not os.path.exists(path_to_temp_table_folder):
            os.mkdir(path_to_temp_table_folder)

    @tricky_loggy
    def create_temp_result_folder(self) -> None:

        _type_of_table = self.proc_obj.type

        _uuid: str = '_'.join((_type_of_table, str(self.proc_obj.folder.unique_identifier)))

        path_to_temp_results_folder = os.path.join(
            os.getcwd(),
            'word',
            'temp_tables',
            'results',
            _uuid,
        )

        if not os.path.exists(path_to_temp_results_folder):
            os.mkdir(path_to_temp_results_folder)

    @tricky_loggy
    def chunk_data_process(self, *args) -> None:

        def copy_table_template(_uuid: uuid, _path_to_copied_file: str) -> None:
            path_to_obj_table = self.template_path

            shutil.copy(path_to_obj_table, _path_to_copied_file)

        _pointer, proc_data, _semaphore = args

        with _semaphore:

            _uuid: uuid = uuid.uuid4()

            _type_of_table: str = self.proc_obj.type
            _is_table: bool = self.proc_obj.settings.get('table')

            temp_table_template_folder_name: str = '_'.join(
                (_type_of_table, str(self.proc_obj.folder.unique_identifier)))
            temp_table_result_folder_name: str = '_'.join((_type_of_table, str(self.proc_obj.folder.unique_identifier)))

            path_to_copied_file: str = os.path.join(
                os.getcwd(),
                'word',
                'temp_tables',
                'templates',
                temp_table_template_folder_name,
                str(_pointer) + '_' + str(_uuid) + '.docx',
            )

            _output_path: str = os.path.join(
                os.getcwd(),
                'word',
                'temp_tables',
                'results',
                temp_table_result_folder_name,
                str(_pointer) + '_' + str(_uuid) + '.docx',
            )

            copy_table_template(_uuid, path_to_copied_file)

            _template: DocxTemplate = DocxTemplate(path_to_copied_file)

            try:
                _template.render(
                    {
                        'columns': proc_data[0].keys(),
                        'data': proc_data,
                        'is_table': _is_table,
                    }, autoescape=True)

            except IndexError:
                raise IndexError('Невозможно сформировавть таблицу по причине нехватки данных!')

            settings: dict = {k: v for (k, v) in self.proc_obj.settings.items()}
            static_settings: dict = self.proc_obj.static_settings

            tags: list = self.proc_obj.response_part.get('query_ar')
            tags_highlight_settings: dict = self.proc_obj.settings.get('tag_highlight')

            if self.proc_obj.settings.get('table'):
                styles = TableStylesGenerator(
                    _template,
                    settings,
                    static_settings,
                    tags,
                    tags_highlight_settings,
                )
                setattr(styles, 'pointer', _pointer)
                styles.apply_table_styles()
            else:
                styles = SchedulerStylesGenerator(
                    _template,
                    settings,
                    static_settings,
                    tags,
                    tags_highlight_settings,
                )
                setattr(styles, 'pointer', _pointer)
                styles.apply_scheduler_styles()

            _template.save(_output_path)

    @tricky_loggy
    def merge_procs_tables(self) -> None:

        _is_table: bool = self.proc_obj.settings.get('table')
        _uuid: str = str(self.proc_obj.folder.unique_identifier)
        _type_of_table: str = self.proc_obj.type

        results_folder_name: str = '_'.join((_type_of_table, _uuid))

        path_to_results: str = os.path.join(
            os.getcwd(),
            'word',
            'temp_tables',
            'results',
            results_folder_name,
        )

        path_to_out_file: str = os.path.join(
            os.getcwd(),
            'word',
            'temp_tables',
            'out.docx',
        )

        def add_heading_to_table(_master: docx.Document) -> None:
            def choose_title(_is_table_: bool, _type_: str) -> str:

                lang_dicts: dict = ReportLanguagePicker(self.report_format)()

                dict_obj = lang_dicts.get('titles')

                def inner(name, __type):
                    obj = dict_obj.get(name)
                    if __type == 'soc':
                        return obj.get('soc')
                    return obj.get('smi')

                if _is_table_:
                    return inner('table', _type_)
                else:
                    return inner('scheduler', _type_)

            title = choose_title(_is_table, _type_of_table)

            style_heading = _master.styles['Intense Quote']

            _paragraph = _master.paragraphs[0]
            p = _paragraph._element
            p.getparent().remove(p)
            p._p = p._element = None

            paragraph = _master.add_paragraph(style=style_heading)
            paragraph.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER

            run = paragraph.add_run(title)
            run.font.size = Pt(16)
            run.font.name = 'Arial'
            run.font.bold = True
            run.font.italic = False
            run.font.color.rgb = RGBColor(19, 104, 181)

        master: docx.Document = docx.Document(path_to_out_file)
        add_heading_to_table(master)
        composer: Composer = Composer(master)

        sorted_list_dir: list = sorted(os.listdir(path_to_results), key=lambda x: int(x.split('_')[0]))

        def delete_paragraph(_paragraph) -> None:
            """ Метод позволяет убрать пустое пространство между таблицами. """

            if self.proc_obj.settings.get('table'):
                p = _paragraph._element
                p.getparent().remove(p)
                p._p = p._element = None

        for idx, file in enumerate(sorted_list_dir):
            doc: docx.Document = docx.Document(
                os.path.join(
                    path_to_results,
                    file,
                )
            )

            for paragraph in doc.paragraphs:
                delete_paragraph(paragraph)

            composer.append(doc)

            _position = self.proc_obj.settings.get('position')

            composer.save(
                os.path.join(
                    os.getcwd(),
                    'word',
                    'temp',
                    str(_uuid),
                    f'output-{_position}-table.docx',
                )
            )

    @tricky_loggy
    def apply(self) -> None:
        self.create_temp_template_folder()
        self.create_temp_result_folder()

        step = 50

        semaphore = Semaphore(min(32, os.cpu_count()))

        processes = []
        for pointer, chunk in enumerate(range(0, len(self.data), step)):
            process = Process(target=self.chunk_data_process,
                              args=(pointer, self.data[chunk:chunk + step], semaphore))
            processes.append(process)
            process.start()

        for process in processes:
            process.join()

        self.merge_procs_tables()


class ContentProcess(AbstractRunnerMixin):
    def __init__(self, proc_object, data, report_format):
        super().__init__(proc_object, data, report_format)
        self._template_path: str = os.path.join(
            os.getcwd(),
            'word',
            'temp_templates',
            f'{self.proc_obj.folder.unique_identifier}',
            'template_parts',
            'table_of_contents.docx',
        )

    @tricky_loggy
    def apply(self) -> None:
        path_to_obj_table_of_contents = self.template_path
        template = DocxTemplate(path_to_obj_table_of_contents)

        position = self.proc_obj.settings.get('position')

        output_path = os.path.join(
            os.getcwd(),
            'word',
            'temp',
            f'{self.proc_obj.folder.unique_identifier}',
            f'output-{position}-content.docx',
        )

        has_soc = self.data.get('soc', {})
        has_smi = self.data.get('smi', {})

        try:
            template.render(
                {
                    'table_of_contents_soc': has_soc,
                    'table_of_contents_smi': has_smi,
                    'lang': self.report_lang,
                }, autoescape=True)
        except IndexError:
            raise IndexError('Невозможно сформировать оглавление по причине нехватки данных!')

        template.save(output_path)


class TagsProcess(AbstractRunnerMixin):
    def __init__(self, proc_object, data, report_format):
        super().__init__(proc_object, data, report_format)
        self._template_path = os.path.join(
            os.getcwd(),
            'word',
            'temp_templates',
            f'{self.proc_obj.folder.unique_identifier}',
            'template_parts',
            'tags.docx',
        )

    @tricky_loggy
    def apply(self) -> None:
        path_to_obj_tags = self.template_path
        template = DocxTemplate(path_to_obj_tags)

        _position = self.proc_obj.settings.get('position')

        output_path = os.path.join(
            os.getcwd(),
            'word',
            'temp',
            f'{self.proc_obj.folder.unique_identifier}',
            f'output-{_position}-atags.docx',
        )

        try:
            template.render({
                'tags': self.data,
                'lang': self.report_lang,
            }, autoescape=True)
        except IndexError:
            raise IndexError('Невозможно сформировать теги по причине нехватки данных!')

        template.save(output_path)


class BaseProcess(AbstractRunnerMixin):
    def __init__(self, proc_object, data, report_format):
        super().__init__(proc_object, data, report_format)
        self._template_path = os.path.join(
            os.getcwd(),
            'word',
            'temp_templates',
            f'{self.proc_obj.folder.unique_identifier}',
            'template_parts',
            'base.docx',
        )

    @tricky_loggy
    def apply(self) -> None:
        position = 0
        path_to_obj_base = self.template_path
        template = DocxTemplate(path_to_obj_base)

        output_path = os.path.join(
            os.getcwd(),
            'word',
            'temp',
            f'{self.proc_obj.folder.unique_identifier}',
            f'output-{position}-base.docx',
        )

        project_name = self.data.get('project_name')
        start_time = self.data.get('start_time')
        end_time = self.data.get('end_time')
        date_of_export = self.data.get('date_of_export')

        template.render({
            'project_name': project_name,
            'date_start': start_time,
            'date_end': end_time,
            'date_of_export': date_of_export,
            'lang': self.report_lang,
        }, autoescape=True)

        template.save(output_path)


class TotalMessagesCountProcess(AbstractRunnerMixin):
    def __init__(self, proc_object, data, report_format):
        super().__init__(proc_object, data, report_format)
        self._template_path = os.path.join(
            os.getcwd(),
            'word',
            'temp_templates',
            f'{self.proc_obj.folder.unique_identifier}',
            'template_parts',
            'total_messages_count.docx',
        )

    @tricky_loggy
    def apply(self) -> None:
        template: DocxTemplate = DocxTemplate(self._template_path)

        _position = self.proc_obj.settings.get('position')

        output_path = os.path.join(
            os.getcwd(),
            'word',
            'temp',
            f'{self.proc_obj.folder.unique_identifier}',
            f'output-{_position}-total-messages-count.docx'
        )

        total_messages_count: int = self.data.get('total_count', 0)
        positive_messages_count: int = self.data.get('pos_count', 0)
        neutral_messages_count: int = self.data.get('neu_count', 0)
        negative_messages_count: int = self.data.get('neg_count', 0)

        messages_count_dict: dict = ReportLanguagePicker(self.report_format)().get('messages_count')

        title = messages_count_dict.get('title')
        total_title = messages_count_dict.get('messages')
        positive_title = messages_count_dict.get('positive')
        neutral_title = messages_count_dict.get('neutral')
        negative_title = messages_count_dict.get('negative')

        context = {
            'title': title,
            'total': total_messages_count,
            'total_title': total_title,
            'neutral': neutral_messages_count,
            'neutral_title': neutral_title,
            'positive': positive_messages_count,
            'positive_title': positive_title,
            'negative': negative_messages_count,
            'negative_title': negative_title
        }

        template.render(context, autoescape=True)

        template.save(output_path)


class MessagesDynamicsProcess(AbstractRunnerMixin):
    def __init__(self, proc_object, data, report_format):
        super().__init__(proc_object, data, report_format)
        self._template_path = os.path.join(
            os.getcwd(),
            'word',
            'temp_templates',
            f'{self.proc_obj.folder.unique_identifier}',
            'template_parts',
            'highcharts',
            'messages_dynamics.docx',
        )

    @tricky_loggy
    def apply(self) -> None:
        """ Не оборачиваю данную функцию в декоратор render_diagram по причине того, что данный процесс использует
        всего один вид графика - линейный и к тому же в клиентских настройках ничего не приходит. """

        template: DocxTemplate = DocxTemplate(self._template_path)

        _position: str = self.proc_obj.settings.get('position')

        messages_dynamic = HighchartsCreator(
            self.report_format,
            self.proc_obj.folder,
        )

        chart = MetricsGenerator()

        start_date = self.proc_obj.response_part.get('s_date')
        end_date = self.proc_obj.response_part.get('f_date')

        soc_news = self.proc_obj.response_part.get('f_news2')
        smi_news = self.proc_obj.response_part.get('f_news')

        categories = chart.define_timedelta(start_date, end_date)
        chart_series = chart.count_messages(soc_news, smi_news)
        query_string: str = messages_dynamic.linear(
            chart_categories=categories,
            chart_series=chart_series,
        )

        class_name = self.__class__.__name__

        path_to_image = os.path.join(
            os.getcwd(),
            'word',
            'highcharts_temp_images',
            f'{self.proc_obj.folder.unique_identifier}',
            class_name + '.png'
        )

        output_path = os.path.join(
            os.getcwd(),
            'word',
            'temp',
            f'{self.proc_obj.folder.unique_identifier}',
            f'output-{_position}-messages-{class_name}.docx'
        )

        response = messages_dynamic.do_post_request_to_highcharts_server(query_string)

        messages_dynamic.save_data_as_png(response, path_to_image)

        dynamics_image = InlineImage(template, image_descriptor=path_to_image, width=Cm(23), height=Cm(13))

        title: str = ReportLanguagePicker(self.report_format)().get('titles').get('messages_dynamics')

        template.render({'title': title, 'image': dynamics_image}, autoescape=True)

        template.save(output_path)


class SentimentsProcess(AbstractRunnerMixin):
    def __init__(self, proc_object, data, report_format):
        super().__init__(proc_object, data, report_format)
        self._template_path = os.path.join(
            os.getcwd(),
            'word',
            'temp_templates',
            f'{self.proc_obj.folder.unique_identifier}',
            'template_parts',
            'highcharts',
            'sentiments.docx',
        )

    @tricky_loggy
    @render_diagram(color_flag='sentiments')
    def apply(self, **kwargs) -> tuple:
        news_count: list = self.proc_obj.response_part.get('news_counts')
        diagram_type: str = kwargs.pop('diagram_type')
        news_count_union: dict = {}

        for n_c in news_count:
            news_count_union.update(n_c)

        title = ReportLanguagePicker(self.report_format)().get('titles').get('sentiments')
        sentiments_translate = ReportLanguagePicker(self.report_format)().get('sentiments_translate')
        positive_title = sentiments_translate.get('positive')
        negative_title = sentiments_translate.get('negative')
        neutral_title = sentiments_translate.get('neutral')

        positive_count = news_count_union.get('pos')
        negative_count = news_count_union.get('neg')
        neutral_count = news_count_union.get('neu')

        sentiments_count = MetricsGenerator().count_percentage_of_sentiments(positive_count, negative_count,
                                                                             neutral_count)

        positive_percent = sentiments_count.get('pos', 0)
        negative_percent = sentiments_count.get('neg', 0)
        neutral_percent = sentiments_count.get('neu', 0)

        chart_categories: list[dict] = generate_chart_categories(
            (positive_title, positive_count), (negative_title, negative_count), (neutral_title, neutral_count)
        )

        chart_series: list[dict] = [
            {
                'type': diagram_type,
                'data': chart_categories,
            },
        ]

        return chart_categories, chart_series, {
            'title': title,
            'positive_title': positive_title,
            'negative_title': negative_title,
            'neutral_title': neutral_title,
            'positive_count': positive_count,
            'negative_count': negative_count,
            'neutral_count': neutral_count,
            'positive_percent': positive_percent,
            'negative_percent': negative_percent,
            'neutral_percent': neutral_percent
        }


class DistributionProcess(AbstractRunnerMixin):
    def __init__(self, proc_object, data, report_format):
        super().__init__(proc_object, data, report_format)
        self._template_path = os.path.join(
            os.getcwd(),
            'word',
            'temp_templates',
            f'{self.proc_obj.folder.unique_identifier}',
            'template_parts',
            'highcharts',
            'distribution.docx',
        )

    @tricky_loggy
    @render_diagram(color_flag='distribution')
    def apply(self, **kwargs) -> tuple:
        soc_count: int = int(self.proc_obj.response_part.get('soc_count'))
        smi_count: int = int(self.proc_obj.response_part.get('smi_count'))
        langs_dict: dict = ReportLanguagePicker(self.report_format)().get('titles')
        title: str = langs_dict.get('distribution')
        soc_title: str = langs_dict.get('soc')
        smi_title: str = langs_dict.get('smi')
        diagram_type: str = kwargs.pop('diagram_type')

        distribution_percents: dict = MetricsGenerator().count_percentage_of_distribution(smi_count, soc_count)

        chart_categories: list[dict] = generate_chart_categories(
            (soc_title, soc_count), (smi_title, smi_count)
        )

        chart_series: list[dict] = [
            {
                'type': diagram_type,
                'data': chart_categories,
            },
        ]

        return chart_categories, chart_series, {
            'title': title,
            'smi_title': smi_title,
            'soc_title': soc_title,
            'smi_count': smi_count,
            'soc_count': soc_count,
            'smi_percent': distribution_percents.get('smi'),
            'soc_percent': distribution_percents.get('soc'),
        }


class SmiDistributionProcess(AbstractRunnerMixin):
    def __init__(self, proc_object, data, report_format):
        super().__init__(proc_object, data, report_format)
        self._template_path = os.path.join(
            os.getcwd(),
            'word',
            'temp_templates',
            f'{self.proc_obj.folder.unique_identifier}',
            'template_parts',
            'highcharts',
            'smi_distribution.docx',
        )

    @tricky_loggy
    @render_diagram(color_flag='smi_distribution', context_flag=True)
    @throw_params_for_distribution_diagram(
        category_names_key='categoryNames',
        title_key='smi_distribution',
        distribution_keys=('name_cat', 'COUNTER'),
        distribution_translate=True,
    )
    def apply(self) -> tuple:
        pass


class SocDistributionProcess(AbstractRunnerMixin):
    def __init__(self, proc_object, data, report_format):
        super().__init__(proc_object, data, report_format)
        self._template_path = os.path.join(
            os.getcwd(),
            'word',
            'temp_templates',
            f'{self.proc_obj.folder.unique_identifier}',
            'template_parts',
            'highcharts',
            'soc_distribution.docx',
        )

    @tricky_loggy
    @render_diagram(color_flag='soc_distribution', context_flag=True)
    @throw_params_for_distribution_diagram(
        category_names_key='categoryNames2',
        title_key='soc_distribution',
        distribution_keys=('name', 'COUNTER'),
    )
    def apply(self) -> tuple:
        pass


class TopMediaProcess(AbstractRunnerMixin):
    def __init__(self, proc_object, data, report_format):
        super().__init__(proc_object, data, report_format)
        self._template_path = os.path.join(
            os.getcwd(),
            'word',
            'temp_templates',
            f'{self.proc_obj.folder.unique_identifier}',
            'template_parts',
            'highcharts',
            'media_top.docx',
        )

    @tricky_loggy
    @render_diagram(color_flag='media_top', context_flag=True)
    @throw_params_for_distribution_diagram(
        category_names_key='itemsCount',
        title_key='media_top',
        distribution_keys=('RESOURCE_NAME', 'COUNTER'),
    )
    def apply(self) -> dict:
        pass


class TopSocialProcess(AbstractRunnerMixin):
    def __init__(self, proc_object, data, report_format):
        super().__init__(proc_object, data, report_format)
        self._template_path = os.path.join(
            os.getcwd(),
            'word',
            'temp_templates',
            f'{self.proc_obj.folder.unique_identifier}',
            'template_parts',
            'highcharts',
            'soc_top.docx',
        )

    @tricky_loggy
    @render_diagram(color_flag='soc_top', context_flag=True)
    @throw_params_for_distribution_diagram(
        category_names_key='itemsCount2',
        title_key='soc_top',
        distribution_keys=('resource_name', 'COUNTER'),
    )
    def apply(self) -> tuple:
        pass


class MostPopularSocProcess(AbstractRunnerMixin):
    def __init__(self, proc_object, data, report_format):
        super().__init__(proc_object, data, report_format)
        self._template_path = os.path.join(
            os.getcwd(),
            'word',
            'temp_templates',
            f'{self.proc_obj.folder.unique_identifier}',
            'template_parts',
            'highcharts',
            'most_popular_soc.docx',
        )

    @tricky_loggy
    @render_diagram(color_flag='most_popular_soc', context_flag=True)
    @throw_params_for_distribution_diagram(
        title_key='most_popular_soc',
        distribution_keys=('resource_name', 'counter'),
        has_distribution='count_most_popular_metrix',
    )
    def apply(self) -> tuple:
        pass


class TopNegativeProcess(AbstractRunnerMixin):
    def __init__(self, proc_object, data, report_format):
        super().__init__(proc_object, data, report_format)
        self._template_path = os.path.join(
            os.getcwd(),
            'word',
            'temp_templates',
            f'{self.proc_obj.folder.unique_identifier}',
            'template_parts',
            'highcharts',
            'top_negative.docx',
        )

    @tricky_loggy
    @render_diagram(color_flag='top_negative', context_flag=True)
    @throw_params_for_distribution_diagram(
        title_key='top_negative',
        distribution_keys=('resource_name', 'counter'),
        has_distribution='count_top_negative',
    )
    def apply(self) -> tuple:
        pass


class SmiTopNegativeProcess(AbstractRunnerMixin):
    def __init__(self, proc_object, data, report_format):
        super().__init__(proc_object, data, report_format)
        self._template_path = os.path.join(
            os.getcwd(),
            'word',
            'temp_templates',
            f'{self.proc_obj.folder.unique_identifier}',
            'template_parts',
            'highcharts',
            'smi_top_negative.docx',
        )

    @tricky_loggy
    @render_diagram(color_flag='smi_top_negative', context_flag=True)
    @throw_params_for_distribution_diagram(
        title_key='smi_top_negative',
        distribution_keys=('resource_name', 'counter'),
        has_distribution='count_top_negative_smi',
    )
    def apply(self) -> tuple | None:
        pass


class SocTopNegativeProcess(AbstractRunnerMixin):
    def __init__(self, proc_object, data, report_format):
        super().__init__(proc_object, data, report_format)
        self._template_path = os.path.join(
            os.getcwd(),
            'word',
            'temp_templates',
            f'{self.proc_obj.folder.unique_identifier}',
            'template_parts',
            'highcharts',
            'soc_top_negative.docx',
        )

    @tricky_loggy
    @render_diagram(color_flag='soc_top_negative', context_flag=True)
    @throw_params_for_distribution_diagram(
        title_key='soc_top_negative',
        distribution_keys=('resource_name', 'counter'),
        has_distribution='count_top_negative_soc',
    )
    def apply(self) -> tuple | None:
        pass


class WorldMapProcess(AbstractRunnerMixin):
    def __init__(self, proc_object, data, report_format):
        super().__init__(proc_object, data, report_format)
        self._template_path = os.path.join(
            os.getcwd(),
            'word',
            'temp_templates',
            f'{self.proc_obj.folder.unique_identifier}',
            'template_parts',
            'highcharts',
            'world_map.docx',
        )

    @tricky_loggy
    @render_map(
        json_marking_title='world_map.JSON',
        region_key='countries_hc',
        stat_map_key='stat_map',
        map_type='world_map',
    )
    def apply(self):
        pass


class KazakhstanMapProcess(AbstractRunnerMixin):
    def __init__(self, proc_object, data, report_format):
        super().__init__(proc_object, data, report_format)
        self._template_path = os.path.join(
            os.getcwd(),
            'word',
            'temp_templates',
            f'{self.proc_obj.folder.unique_identifier}',
            'template_parts',
            'highcharts',
            'kaz_map.docx',
        )

    @tricky_loggy
    @render_map(
        json_marking_title='kazakhstan_map.JSON',
        region_key=None,
        stat_map_key='stat_map2',
        map_type='kaz_map',
    )
    def apply(self) -> tuple | None:
        pass
