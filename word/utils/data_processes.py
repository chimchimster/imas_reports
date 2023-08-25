import os
import uuid
import docx
import shutil

from typing import Any
from docx.shared import Pt, RGBColor
from docxcompose.composer import Composer
from word.local import ReportLanguagePicker
from docxtpl import DocxTemplate, InlineImage
from multiprocessing import Process, Semaphore
from ..tools import TableStylesGenerator, SchedulerStylesGenerator


class ProcessDataGenerator(Process):

    def __init__(self, proc_obj: Any):
        super().__init__()
        self.proc_obj = proc_obj
        self.templates: dict = {
            'table':
                os.path.join(
                    os.getcwd(),
                    'word',
                    'temp_templates',
                    f'{self.proc_obj.folder.unique_identifier}',
                    'template_parts',
                    'table.docx',
                ),
            'table_of_contents':
                os.path.join(
                    os.getcwd(),
                    'word',
                    'temp_templates',
                    f'{self.proc_obj.folder.unique_identifier}',
                    'template_parts',
                    'table_of_contents.docx',
                ),
            'tags':
                os.path.join(
                    os.getcwd(),
                    'word',
                    'temp_templates',
                    f'{self.proc_obj.folder.unique_identifier}',
                    'template_parts',
                    'tags.docx',
                ),
            'base': os.path.join(
                os.getcwd(),
                'word',
                'temp_templates',
                f'{self.proc_obj.folder.unique_identifier}',
                'template_parts',
                'base.docx',
            ),
        }

    def run(self) -> None:

        self.proc_obj.generate_data()

        data = self.proc_obj.data_collection

        report_format = self.proc_obj.static_settings.get('format', 'word_rus')
        report_lang = report_format.split('_')[1]  # TODO: от этого "ужаса" нужно избавиться на уровне клиентской части

        if self.proc_obj.flag == 'table':

            def create_temp_template_folder() -> None:

                _type_of_table = self.proc_obj.type

                _uuid: str = '_'.join((_type_of_table, str(self.proc_obj.folder.unique_identifier)))

                path_to_temp_table_folder = os.path.join(
                    os.getcwd(),
                    'word',
                    'temp_tables',
                    'templates',
                    _uuid,
                )

                if not os.path.exists(path_to_temp_table_folder):
                    os.mkdir(path_to_temp_table_folder)

            def create_temp_result_folder() -> None:

                _type_of_table = self.proc_obj.type

                _uuid: str = '_'.join((_type_of_table, str(self.proc_obj.folder.unique_identifier)))

                path_to_temp_results_folder = os.path.join(
                    os.getcwd(),
                    'word',
                    'temp_tables',
                    'results',
                    _uuid,
                )

                if not os.path.exists(path_to_temp_results_folder):
                    os.mkdir(path_to_temp_results_folder)

            def chunk_data_process(*args) -> None:

                def copy_table_template(_uuid: uuid, _path_to_copied_file: str) -> None:
                    path_to_obj_table = self.templates.get('table')

                    shutil.copy(path_to_obj_table, _path_to_copied_file)

                _pointer, proc_data, _semaphore = args

                with _semaphore:

                    _uuid: uuid = uuid.uuid4()

                    _type_of_table: str = self.proc_obj.type
                    _is_table: bool = self.proc_obj.settings.get('table')

                    temp_table_template_folder_name: str = '_'.join(
                        (_type_of_table, str(self.proc_obj.folder.unique_identifier)))
                    temp_table_result_folder_name: str = '_'.join(
                        (_type_of_table, str(self.proc_obj.folder.unique_identifier)))

                    path_to_copied_file: str = os.path.join(
                        os.getcwd(),
                        'word',
                        'temp_tables',
                        'templates',
                        temp_table_template_folder_name,
                        str(pointer) + '_' + str(_uuid) + '.docx',
                    )

                    _output_path: str = os.path.join(
                        os.getcwd(),
                        'word',
                        'temp_tables',
                        'results',
                        temp_table_result_folder_name,
                        str(pointer) + '_' + str(_uuid) + '.docx',
                    )

                    copy_table_template(_uuid, path_to_copied_file)

                    _template: DocxTemplate = DocxTemplate(path_to_copied_file)

                    try:
                        _template.render(
                            {
                                'columns': proc_data[0].keys(),
                                'data': proc_data,
                                'is_table': _is_table,
                            }, autoescape=True)

                    except IndexError:
                        raise IndexError('Невозможно сформировавть таблицу по причине нехватки данных!')

                    settings: dict = {k: v for (k, v) in self.proc_obj.settings.items()}
                    static_settings: dict = self.proc_obj.static_settings

                    tags: list = self.proc_obj.response_part.get('query_ar')
                    tags_highlight_settings: dict = self.proc_obj.settings.get('tag_highlight')

                    if self.proc_obj.settings.get('table'):
                        styles = TableStylesGenerator(
                            _template,
                            settings,
                            static_settings,
                            tags,
                            tags_highlight_settings,
                        )
                        setattr(styles, 'pointer', pointer)
                        styles.apply_table_styles()
                    else:
                        styles = SchedulerStylesGenerator(
                            _template,
                            settings,
                            static_settings,
                            tags,
                            tags_highlight_settings,
                        )
                        setattr(styles, 'pointer', pointer)
                        styles.apply_scheduler_styles()

                    _template.save(_output_path)

            def merge_procs_tables() -> None:

                _is_table: bool = self.proc_obj.settings.get('table')
                _uuid: str = str(self.proc_obj.folder.unique_identifier)
                _type_of_table: str = self.proc_obj.type

                results_folder_name: str = '_'.join((_type_of_table, _uuid))

                path_to_results: str = os.path.join(
                    os.getcwd(),
                    'word',
                    'temp_tables',
                    'results',
                    results_folder_name,
                )

                path_to_out_file: str = os.path.join(
                    os.getcwd(),
                    'word',
                    'temp_tables',
                    'out.docx',
                )

                def add_heading_to_table(_master: docx.Document) -> None:
                    def choose_title(_is_table_: bool, _type_: str) -> str:

                        lang_dicts: dict = ReportLanguagePicker(report_format)()

                        dict_obj = lang_dicts.get('titles')

                        def inner(name, __type):
                            obj = dict_obj.get(name)
                            if __type == 'soc':
                                return obj.get('soc')
                            return obj.get('smi')

                        if _is_table_:
                            return inner('table', _type_)
                        else:
                            return inner('scheduler', _type_)

                    title = choose_title(_is_table, _type_of_table)

                    style_heading = _master.styles['Intense Quote']

                    _paragraph = _master.paragraphs[0]
                    p = _paragraph._element
                    p.getparent().remove(p)
                    p._p = p._element = None

                    paragraph = _master.add_paragraph(style=style_heading)
                    paragraph.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER

                    run = paragraph.add_run(title)
                    run.font.size = Pt(20)
                    run.font.name = 'Arial'
                    run.font.bold = False
                    run.font.italic = False
                    run.font.color.rgb = RGBColor(0, 0, 0)

                master: docx.Document = docx.Document(path_to_out_file)
                add_heading_to_table(master)
                composer: Composer = Composer(master)

                sorted_list_dir: list = sorted(os.listdir(path_to_results), key=lambda x: int(x.split('_')[0]))

                for idx, file in enumerate(sorted_list_dir):

                    doc: docx.Document = docx.Document(
                        os.path.join(
                            path_to_results,
                            file,
                        )
                    )

                    def delete_paragraph(_paragraph) -> None:
                        """ Метод позволяет убрать пустое пространство между таблицами. """

                        if self.proc_obj.settings.get('table'):
                            p = _paragraph._element
                            p.getparent().remove(p)
                            p._p = p._element = None

                    for paragraph in doc.paragraphs:
                        delete_paragraph(paragraph)

                    composer.append(doc)

                _position = self.proc_obj.settings.get('position')

                composer.save(
                    os.path.join(
                        os.getcwd(),
                        'word',
                        'temp',
                        str(_uuid),
                        f'output-{_position}-table.docx',
                    )
                )

            create_temp_template_folder()
            create_temp_result_folder()

            max_processes = 50
            step = 50

            semaphore = Semaphore(max_processes)

            processes = []
            for pointer, chunk in enumerate(range(0, len(data), step)):
                process = Process(target=chunk_data_process,
                                  args=(pointer, data[chunk:chunk + step], semaphore))
                processes.append(process)
                process.start()

            for process in processes:
                process.join()

            merge_procs_tables()

        elif self.proc_obj.flag == 'content':

            path_to_obj_table_of_contents = self.templates.get('table_of_contents')
            template = DocxTemplate(path_to_obj_table_of_contents)

            position = self.proc_obj.settings.get('position')

            output_path = os.path.join(
                os.getcwd(),
                'word',
                'temp',
                f'{self.proc_obj.folder.unique_identifier}',
                f'output-{position}-content.docx',
            )

            has_soc = data.get('soc', {})
            has_smi = data.get('smi', {})

            try:
                template.render(
                    {
                        'table_of_contents_soc': has_soc,
                        'table_of_contents_smi': has_smi,
                        'lang': report_lang,
                    }, autoescape=True)
            except IndexError:
                raise IndexError('Невозможно сформировать оглавление по причине нехватки данных!')

            template.save(output_path)

        elif self.proc_obj.flag == 'tags':

            path_to_obj_tags = self.templates.get('tags')
            template = DocxTemplate(path_to_obj_tags)

            _position = self.proc_obj.settings.get('position')

            output_path = os.path.join(
                os.getcwd(),
                'word',
                'temp',
                f'{self.proc_obj.folder.unique_identifier}',
                f'output-{_position}-atags.docx',
            )

            try:
                template.render({
                    'tags': data,
                    'lang': report_lang,
                }, autoescape=True)
            except IndexError:
                raise IndexError('Невозможно сформировать теги по причине нехватки данных!')
            template.save(output_path)

        elif self.proc_obj.flag == 'base':

            position = 0
            path_to_obj_base = self.templates.get('base')
            template = DocxTemplate(path_to_obj_base)

            output_path = os.path.join(
                os.getcwd(),
                'word',
                'temp',
                f'{self.proc_obj.folder.unique_identifier}',
                f'output-{position}-base.docx',
            )

            project_name = data.get('project_name')
            start_time = data.get('start_time')
            end_time = data.get('end_time')
            date_of_export = data.get('date_of_export')

            template.render({
                'project_name': project_name,
                'date_start': start_time,
                'date_end': end_time,
                'date_of_export': date_of_export,
                'lang': report_lang,
            }, autoescape=True)

            template.save(output_path)
