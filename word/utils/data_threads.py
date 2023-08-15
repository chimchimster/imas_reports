import multiprocessing
import os
import uuid
from concurrent.futures import ProcessPoolExecutor

import docx
import shutil


from typing import Any
from docxcompose.composer import Composer
from multiprocessing import Pool, Process, Barrier, Lock, Semaphore, Queue
from docxtpl import DocxTemplate, InlineImage
from ..tools import TableStylesGenerator, SchedulerStylesGenerator


class ProcessDataGenerator(Process):

    def __init__(self, proc_obj: Any):
        super().__init__()
        self.proc_obj = proc_obj
        self.templates: dict = {
            'table':
                os.path.join(
                    os.getcwd(),
                    'word',
                    'temp_templates',
                    f'{self.proc_obj.folder.unique_identifier}',
                    'template_parts',
                    'table.docx'
                ),
            'table_of_contents':
                os.path.join(
                    os.getcwd(),
                    'word',
                    'temp_templates',
                    f'{self.proc_obj.folder.unique_identifier}',
                    'template_parts',
                    'table_of_contents.docx'
                ),
            'tags':
                os.path.join(
                    os.getcwd(),
                    'word',
                    'temp_templates',
                    f'{self.proc_obj.folder.unique_identifier}',
                    'template_parts',
                    'tags.docx'
                ),
            'base': os.path.join(
                os.getcwd(),
                'word',
                'temp_templates',
                f'{self.proc_obj.folder.unique_identifier}',
                'template_parts',
                'base.docx'
            ),
        }

    def run(self) -> None:

        self.proc_obj.generate_data()

        data = self.proc_obj.data_collection

        report_format = self.proc_obj._static_rest_data.get('format', 'word_rus')
        report_lang = report_format.split('_')[1]

        if self.proc_obj.flag == 'table':

            def create_temp_template_folder() -> None:

                _type_of_table = self.proc_obj._type

                _uuid: str = '_'.join((_type_of_table, str(self.proc_obj.folder.unique_identifier)))

                path_to_temp_table_folder = os.path.join(
                    os.getcwd(),
                    'word',
                    'temp_tables',
                    'templates',
                    _uuid,
                )

                if not os.path.exists(path_to_temp_table_folder):
                    os.mkdir(path_to_temp_table_folder)

            def create_temp_result_folder() -> None:

                _type_of_table = self.proc_obj._type

                _uuid: str = '_'.join((_type_of_table, str(self.proc_obj.folder.unique_identifier)))

                path_to_temp_results_folder = os.path.join(
                    os.getcwd(),
                    'word',
                    'temp_tables',
                    'results',
                    _uuid,
                )

                if not os.path.exists(path_to_temp_results_folder):
                    os.mkdir(path_to_temp_results_folder)

            def chunk_data_process(*args):

                def copy_table_template(_uuid: uuid, _path_to_copied_file: str) -> None:
                    table_docx_obj = self.templates.get('table')

                    shutil.copy(table_docx_obj, _path_to_copied_file)

                _pointer, proc_data, _semaphore = args

                with _semaphore:

                    _uuid: uuid = uuid.uuid4()

                    _type_of_table: str = self.proc_obj._type

                    temp_table_template_folder_name: str = '_'.join((_type_of_table, str(self.proc_obj.folder.unique_identifier)))
                    temp_table_result_folder_name: str = '_'.join((_type_of_table, str(self.proc_obj.folder.unique_identifier)))

                    path_to_copied_file: str = os.path.join(
                        os.getcwd(),
                        'word',
                        'temp_tables',
                        'templates',
                        temp_table_template_folder_name,
                        str(pointer) + '_' + str(_uuid) + '.docx',
                    )

                    _output_path: str = os.path.join(
                        os.getcwd(),
                        'word',
                        'temp_tables',
                        'results',
                        temp_table_result_folder_name,
                        str(pointer) + '_' + str(_uuid) + '.docx',
                    )

                    copy_table_template(_uuid, path_to_copied_file)

                    _template: DocxTemplate = DocxTemplate(path_to_copied_file)

                    try:
                        _template.render(
                            {
                                'columns': proc_data[0].keys(),
                                'data': proc_data,
                                'is_table': True,
                            }, autoescape=True)

                    except IndexError:
                        raise IndexError('Невозможно сформировавть таблицу по причине нехватки данных!')

                    # _semaphore.release()

                    settings: dict = {k: v for (k, v) in self.proc_obj._rest_data.items()}

                    tags: str = self.proc_obj._data.get('query_ar')

                    tags_highlight_settings: dict = self.proc_obj._rest_data.get('tag_highlight')

                    static_rest_data: dict = self.proc_obj._static_rest_data

                    if self.proc_obj._rest_data.get('table'):
                        styles = TableStylesGenerator(_template, tags, settings, tags_highlight_settings, static_rest_data)
                        setattr(styles, 'pointer', pointer)
                        styles.apply_table_styles()
                    else:
                        styles = SchedulerStylesGenerator(_template, tags, settings, tags_highlight_settings)
                        setattr(styles, 'pointer', pointer)
                        styles.apply_scheduler_styles()

                    _template.save(_output_path)

            def merge_procs_tables() -> None:

                _uuid: str = str(self.proc_obj.folder.unique_identifier)
                _type_of_table: str = self.proc_obj._type

                results_folder_name: str = '_'.join((_type_of_table, _uuid))

                path_to_results: str = os.path.join(
                    os.getcwd(),
                    'word',
                    'temp_tables',
                    'results',
                    results_folder_name,
                )

                path_to_out_file: str = os.path.join(
                    os.getcwd(),
                    'word',
                    'temp_tables',
                    'out.docx',
                )

                master: docx.Document = docx.Document(path_to_out_file)
                composer: Composer = Composer(master)

                sorted_list_dir: list = sorted(os.listdir(path_to_results), key=lambda x: int(x.split('_')[0]))

                for idx, file in enumerate(sorted_list_dir):

                    doc: docx.Document = docx.Document(
                        os.path.join(
                            path_to_results,
                            file,
                        )
                    )

                    def delete_paragraph(_paragraph):
                        p = _paragraph._element
                        p.getparent().remove(p)
                        p._p = p._element = None

                    for paragraph in doc.paragraphs:
                        delete_paragraph(paragraph)

                    composer.append(doc)

                position = self.proc_obj._rest_data.get('position')

                composer.save(
                    os.path.join(
                        os.getcwd(),
                        'word',
                        'temp',
                        str(_uuid),
                        f'output-{position}-table.docx',
                    )
                )

            create_temp_template_folder()
            create_temp_result_folder()

            max_processes = 25
            step = 50

            semaphore = multiprocessing.Semaphore(max_processes)

            processes = []
            for pointer, chunk in enumerate(range(0, len(data), step)):
                process = multiprocessing.Process(target=chunk_data_process,
                                                  args=(pointer, data[chunk:chunk + step], semaphore))
                processes.append(process)
                process.start()

            for process in processes:
                process.join()

            merge_procs_tables()

        elif self.proc_obj.flag == 'content':

            template = DocxTemplate(self.templates.get('table_of_contents'))

            position = self.proc_obj._rest_data.get('position')

            output_path = os.path.join(
                os.getcwd(),
                'word',
                'temp',
                f'{self.proc_obj.folder.unique_identifier}',
                f'output-{position}-content.docx',
            )

            has_soc = data.get('soc', {})
            has_smi = data.get('smi', {})

            try:
                template.render(
                    {
                        'table_of_contents_soc': has_soc,
                        'table_of_contents_smi': has_smi,
                        'lang': report_lang,
                    }, autoescape=True)
            except IndexError:
                raise IndexError('Невозможно сформировать оглавление по причине нехватки данных!')

            template.save(output_path)

        elif self.proc_obj.flag == 'tags':

            template = DocxTemplate(self.templates.get('tags'))

            position = self.proc_obj._rest_data.get('position')

            output_path = os.path.join(
                os.getcwd(),
                'word',
                'temp',
                f'{self.proc_obj.folder.unique_identifier}',
                f'output-{position}-atags.docx',
            )

            try:
                template.render({
                    'tags': data,
                    'lang': report_lang,
                }, autoescape=True)
            except IndexError:
                raise IndexError('Невозможно сформировать теги по причине нехватки данных!')
            template.save(output_path)

        elif self.proc_obj.flag == 'base':

            position = 0

            template = DocxTemplate(self.templates.get('base'))

            output_path = os.path.join(
                os.getcwd(),
                'word',
                'temp',
                f'{self.proc_obj.folder.unique_identifier}',
                f'output-{position}-base.docx',
            )

            project_name = data.get('project_name')
            start_time = data.get('start_time')
            end_time = data.get('end_time')
            date_of_export = data.get('date_of_export')

            template.render({
                'project_name': project_name,
                'date_start': start_time,
                'date_end': end_time,
                'date_of_export': date_of_export,
                'lang': report_lang,
            }, autoescape=True)

            template.save(output_path)
