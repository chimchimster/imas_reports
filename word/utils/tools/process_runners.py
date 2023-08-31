import os
import uuid
import docx
import shutil

from docxtpl import DocxTemplate
from docx.shared import RGBColor, Pt
from docxcompose.composer import Composer

from multiprocessing import Semaphore, Process

from .mixins import PropertyProcessesMixin, AbstractRunnerMixin

from word.local import ReportLanguagePicker
from word.tools import (BasePageDataGenerator, TagsGenerator, ContentGenerator,
                        TableContentGenerator, TableStylesGenerator, SchedulerStylesGenerator)


class TableProcess(AbstractRunnerMixin, PropertyProcessesMixin):

    def __init__(self, proc_object, data, report_format):
        super().__init__(proc_object, data, report_format)
        self._template_path: str = os.path.join(
            os.getcwd(),
            'word',
            'temp_templates',
            f'{self.proc_obj.folder.unique_identifier}',
            'template_parts',
            'table.docx',
        )

    def create_temp_template_folder(self) -> None:
        _type_of_table = self.proc_obj.type

        _uuid: str = '_'.join((_type_of_table, str(self.proc_obj.folder.unique_identifier)))

        path_to_temp_table_folder = os.path.join(
            os.getcwd(),
            'word',
            'temp_tables',
            'templates',
            _uuid,
        )

        if not os.path.exists(path_to_temp_table_folder):
            os.mkdir(path_to_temp_table_folder)

    def create_temp_result_folder(self) -> None:

        _type_of_table = self.proc_obj.type

        _uuid: str = '_'.join((_type_of_table, str(self.proc_obj.folder.unique_identifier)))

        path_to_temp_results_folder = os.path.join(
            os.getcwd(),
            'word',
            'temp_tables',
            'results',
            _uuid,
        )

        if not os.path.exists(path_to_temp_results_folder):
            os.mkdir(path_to_temp_results_folder)

    def chunk_data_process(self, *args) -> None:

        def copy_table_template(_uuid: uuid, _path_to_copied_file: str) -> None:
            path_to_obj_table = self.template_path

            shutil.copy(path_to_obj_table, _path_to_copied_file)

        _pointer, proc_data, _semaphore = args

        with _semaphore:

            _uuid: uuid = uuid.uuid4()

            _type_of_table: str = self.proc_obj.type
            _is_table: bool = self.proc_obj.settings.get('table')

            temp_table_template_folder_name: str = '_'.join((_type_of_table, str(self.proc_obj.folder.unique_identifier)))
            temp_table_result_folder_name: str = '_'.join((_type_of_table, str(self.proc_obj.folder.unique_identifier)))

            path_to_copied_file: str = os.path.join(
                os.getcwd(),
                'word',
                'temp_tables',
                'templates',
                temp_table_template_folder_name,
                str(_pointer) + '_' + str(_uuid) + '.docx',
            )

            _output_path: str = os.path.join(
                os.getcwd(),
                'word',
                'temp_tables',
                'results',
                temp_table_result_folder_name,
                str(_pointer) + '_' + str(_uuid) + '.docx',
            )

            copy_table_template(_uuid, path_to_copied_file)

            _template: DocxTemplate = DocxTemplate(path_to_copied_file)

            try:
                _template.render(
                    {
                        'columns': proc_data[0].keys(),
                        'data': proc_data,
                        'is_table': _is_table,
                    }, autoescape=True)

            except IndexError:
                raise IndexError('Невозможно сформировавть таблицу по причине нехватки данных!')

            settings: dict = {k: v for (k, v) in self.proc_obj.settings.items()}
            static_settings: dict = self.proc_obj.static_settings

            tags: list = self.proc_obj.response_part.get('query_ar')
            tags_highlight_settings: dict = self.proc_obj.settings.get('tag_highlight')

            if self.proc_obj.settings.get('table'):
                styles = TableStylesGenerator(
                    _template,
                    settings,
                    static_settings,
                    tags,
                    tags_highlight_settings,
                )
                setattr(styles, 'pointer', _pointer)
                styles.apply_table_styles()
            else:
                styles = SchedulerStylesGenerator(
                    _template,
                    settings,
                    static_settings,
                    tags,
                    tags_highlight_settings,
                )
                setattr(styles, 'pointer', _pointer)
                styles.apply_scheduler_styles()

            _template.save(_output_path)

    def merge_procs_tables(self) -> None:

        _is_table: bool = self.proc_obj.settings.get('table')
        _uuid: str = str(self.proc_obj.folder.unique_identifier)
        _type_of_table: str = self.proc_obj.type

        results_folder_name: str = '_'.join((_type_of_table, _uuid))

        path_to_results: str = os.path.join(
            os.getcwd(),
            'word',
            'temp_tables',
            'results',
            results_folder_name,
        )

        path_to_out_file: str = os.path.join(
            os.getcwd(),
            'word',
            'temp_tables',
            'out.docx',
        )

        def add_heading_to_table(_master: docx.Document) -> None:
            def choose_title(_is_table_: bool, _type_: str) -> str:

                lang_dicts: dict = ReportLanguagePicker(self.report_format)()

                dict_obj = lang_dicts.get('titles')

                def inner(name, __type):
                    obj = dict_obj.get(name)
                    if __type == 'soc':
                        return obj.get('soc')
                    return obj.get('smi')

                if _is_table_:
                    return inner('table', _type_)
                else:
                    return inner('scheduler', _type_)

            title = choose_title(_is_table, _type_of_table)

            style_heading = _master.styles['Intense Quote']

            _paragraph = _master.paragraphs[0]
            p = _paragraph._element
            p.getparent().remove(p)
            p._p = p._element = None

            paragraph = _master.add_paragraph(style=style_heading)
            paragraph.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER

            run = paragraph.add_run(title)
            run.font.size = Pt(20)
            run.font.name = 'Arial'
            run.font.bold = False
            run.font.italic = False
            run.font.color.rgb = RGBColor(0, 0, 0)

        master: docx.Document = docx.Document(path_to_out_file)
        add_heading_to_table(master)
        composer: Composer = Composer(master)

        sorted_list_dir: list = sorted(os.listdir(path_to_results), key=lambda x: int(x.split('_')[0]))

        def delete_paragraph(_paragraph) -> None:
            """ Метод позволяет убрать пустое пространство между таблицами. """

            if self.proc_obj.settings.get('table'):
                p = _paragraph._element
                p.getparent().remove(p)
                p._p = p._element = None

        for idx, file in enumerate(sorted_list_dir):
            doc: docx.Document = docx.Document(
                os.path.join(
                    path_to_results,
                    file,
                )
            )

            for paragraph in doc.paragraphs:
                delete_paragraph(paragraph)

            composer.append(doc)

            _position = self.proc_obj.settings.get('position')

            composer.save(
                os.path.join(
                    os.getcwd(),
                    'word',
                    'temp',
                    str(_uuid),
                    f'output-{_position}-table.docx',
                )
            )

    def apply(self) -> None:
        self.create_temp_template_folder()
        self.create_temp_result_folder()

        max_processes = 50
        step = 50

        semaphore = Semaphore(max_processes)

        processes = []
        for pointer, chunk in enumerate(range(0, len(self.data), step)):
            process = Process(target=self.chunk_data_process,
                              args=(pointer, self.data[chunk:chunk + step], semaphore))
            processes.append(process)
            process.start()

        for process in processes:
            process.join()

        self.merge_procs_tables()


class ContentProcess(AbstractRunnerMixin, PropertyProcessesMixin):
    def __init__(self, proc_object, data, report_format):
        super().__init__(proc_object, data, report_format)
        self._template_path: str = os.path.join(
            os.getcwd(),
            'word',
            'temp_templates',
            f'{self.proc_obj.folder.unique_identifier}',
            'template_parts',
            'table_of_contents.docx',
        )

    def apply(self) -> None:
        path_to_obj_table_of_contents = self.template_path
        template = DocxTemplate(path_to_obj_table_of_contents)

        position = self.proc_obj.settings.get('position')

        output_path = os.path.join(
            os.getcwd(),
            'word',
            'temp',
            f'{self.proc_obj.folder.unique_identifier}',
            f'output-{position}-content.docx',
        )

        has_soc = self.data.get('soc', {})
        has_smi = self.data.get('smi', {})

        try:
            template.render(
                {
                    'table_of_contents_soc': has_soc,
                    'table_of_contents_smi': has_smi,
                    'lang': self.report_lang,
                }, autoescape=True)
        except IndexError:
            raise IndexError('Невозможно сформировать оглавление по причине нехватки данных!')

        template.save(output_path)


class TagsProcess(AbstractRunnerMixin, PropertyProcessesMixin):
    def __init__(self, proc_object, data, report_format):
        super().__init__(proc_object, data, report_format)
        self._template_path = os.path.join(
                    os.getcwd(),
                    'word',
                    'temp_templates',
                    f'{self.proc_obj.folder.unique_identifier}',
                    'template_parts',
                    'tags.docx',
                )

    def apply(self) -> None:
        path_to_obj_tags = self.template_path
        template = DocxTemplate(path_to_obj_tags)

        _position = self.proc_obj.settings.get('position')

        output_path = os.path.join(
            os.getcwd(),
            'word',
            'temp',
            f'{self.proc_obj.folder.unique_identifier}',
            f'output-{_position}-atags.docx',
        )

        try:
            template.render({
                'tags': self.data,
                'lang': self.report_lang,
            }, autoescape=True)
        except IndexError:
            raise IndexError('Невозможно сформировать теги по причине нехватки данных!')

        template.save(output_path)


class BaseProcess(AbstractRunnerMixin, PropertyProcessesMixin):
    def __init__(self, proc_object, data, report_format):
        super().__init__(proc_object, data, report_format)
        self._template_path = os.path.join(
                os.getcwd(),
                'word',
                'temp_templates',
                f'{self.proc_obj.folder.unique_identifier}',
                'template_parts',
                'base.docx',
            )

    def apply(self) -> None:
        position = 0
        path_to_obj_base = self.template_path
        template = DocxTemplate(path_to_obj_base)

        output_path = os.path.join(
            os.getcwd(),
            'word',
            'temp',
            f'{self.proc_obj.folder.unique_identifier}',
            f'output-{position}-base.docx',
        )

        project_name = self.data.get('project_name')
        start_time = self.data.get('start_time')
        end_time = self.data.get('end_time')
        date_of_export = self.data.get('date_of_export')

        template.render({
            'project_name': project_name,
            'date_start': start_time,
            'date_end': end_time,
            'date_of_export': date_of_export,
            'lang': self.report_lang,
        }, autoescape=True)

        template.save(output_path)


class TotalMessagesCountProcess(AbstractRunnerMixin, PropertyProcessesMixin):
    def __init__(self, proc_object, data, report_format):
        super().__init__(proc_object, data, report_format)
        self._template_path = os.path.join(
            os.getcwd(),
            'word',
            'temp_templates',
            f'{self.proc_obj.folder.unique_identifier}',
            'template_parts',
            'total_messages_count.docx',
        )

    def apply(self) -> None:

        template: DocxTemplate = DocxTemplate(self._template_path)

        _position = self.proc_obj.settings.get('position')

        output_path = os.path.join(
            os.getcwd(),
            'word',
            'temp',
            f'{self.proc_obj.folder.unique_identifier}',
            f'output-{_position}-total-messages-count.docx'
        )

        total_messages_count: int = self.data.get('total_count', 0)
        positive_messages_count: int = self.data.get('pos_count', 0)
        neutral_messages_count: int = self.data.get('neu_count', 0)
        negative_messages_count: int = self.data.get('neg_count', 0)

        messages_count_dict: dict = ReportLanguagePicker(self.report_format)().get('messages_count')

        title = messages_count_dict.get('title')
        total_title = messages_count_dict.get('messages')
        positive_title = messages_count_dict.get('positive')
        neutral_title = messages_count_dict.get('neutral')
        negative_title = messages_count_dict.get('negative')

        context = {
            'title': title,
            'total': total_messages_count,
            'total_title': total_title,
            'neutral': neutral_messages_count,
            'neutral_title': neutral_title,
            'positive': positive_messages_count,
            'positive_title': positive_title,
            'negative': negative_messages_count,
            'negative_title': negative_title
        }

        template.render(context, autoescape=True)

        template.save(output_path)