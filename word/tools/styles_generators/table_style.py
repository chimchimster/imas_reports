import re
import docx

from docxtpl import DocxTemplate
from docx.oxml.ns import nsdecls, qn
from docx.shared import RGBColor, Pt, Cm

from logs.decorators import tricky_loggy
from word.local import ReportLanguagePicker
from docx.oxml import parse_xml, OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from ..data_generators.table_data import TableContentGenerator


class TableStylesGenerator(TableContentGenerator):

    translator_smi = {}
    translator_soc = {}

    def __init__(
            self,
            template: DocxTemplate,
            settings: dict,
            static_settings: dict,
            tags: list,
            tags_highlight_settings: dict,
    ):
        self._template = template
        self._settings = settings
        self._static_settings = static_settings
        self._tags = tags
        self._tags_highlight_settings = tags_highlight_settings
        self.pick_language('styles')

    @tricky_loggy
    def apply_table_styles(self):

        table = self.template.tables[0]
        table.style = 'Table Grid'

        if not self.settings:
            return

        if self.settings.get('id') == 'soc':
            self.choose_particular_table_styles(self.translator_soc, table, 'soc')
        else:
            self.choose_particular_table_styles(self.translator_smi, table, 'smi')

        self.delete_first_row(table)

    @tricky_loggy
    def choose_particular_table_styles(self, translator_obj, table_obj, _type):
        def set_cell_width():
            match cell.text:
                case '№':
                    table_obj.columns[idx].width = Cm(1)
                case 'Заголовок' | 'Хабарлама тақырыбы' | 'Title':
                    table_obj.columns[idx].width = Cm(8)
                case "Пост" | 'Краткое содержание' | 'Қысқаша мазмұны' | 'Summary' | 'Post':
                    table_obj.columns[idx].width = Cm(15)
                case 'Дата' | 'Күні' | 'Date':
                    table_obj.columns[idx].width = Cm(5)
                case 'Соцсеть' | 'Категория' | 'Категориясы' | 'Әлеуметтік желілер атауы' | 'Category' | 'Social media':
                    table_obj.columns[idx].width = Cm(5)
                case 'URL':
                    table_obj.columns[idx].width = Cm(5)
                case 'Сообщество' | 'Наименование СМИ' | 'БАҚ атауы' | 'Қауымдастық' | 'Media name' | 'Community':
                    table_obj.columns[idx].width = Cm(8)
                case 'Тональность' | 'Реңкілік' | 'Sentiment':
                    table_obj.columns[idx].width = Cm(6)

        _format = self.static_settings.get('format', 'word_rus')

        dict_languages = ReportLanguagePicker(_format)()

        link_name = dict_languages.get('link', 'Ссылка')

        for column in self.settings['columns']:
            if column.get('id') in translator_obj:
                column_name_en = column.get('id')
                column_name = translator_obj[column_name_en]

                for idx, cell in enumerate(table_obj.row_cells(0)):
                    set_cell_width()

                    if cell.text == column_name:
                        for row in table_obj.rows[1:]:
                            cell = row.cells[idx]

                            self.define_color_of_sentiment_cell(cell)

                            for paragraph in cell.paragraphs:
                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                for run in paragraph.runs:

                                    self.highlight_tag(run, paragraph, self.tags, column_name, self.tags_highlight_settings)

                                    if re.match(r'https?://(?:[-\w.]|(?:%[\da-fA-F]{2}))+[/\w .?=-]*', cell.text) and column_name == 'URL':
                                        hyperlink = self.add_hyperlink(paragraph, cell.text.strip(), link_name, '#0000FF', '#000080')

                                        for old_run in paragraph.runs:
                                            if old_run != hyperlink:
                                                paragraph._p.remove(old_run._r)

                                    self.apply_run_styles(run, column)

    @staticmethod
    def add_hyperlink(paragraph, url, text, color, underline):

        part = paragraph.part
        r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

        new_run = docx.oxml.shared.OxmlElement('w:r')

        rPr = docx.oxml.shared.OxmlElement('w:rPr')

        if not color is None:
            c = docx.oxml.shared.OxmlElement('w:color')
            c.set(docx.oxml.shared.qn('w:val'), color)
            rPr.append(c)

        if not underline:
            u = docx.oxml.shared.OxmlElement('w:u')
            u.set(docx.oxml.shared.qn('w:val'), 'none')
            rPr.append(u)

        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)

        paragraph.add_run(' ')
        paragraph._p.append(hyperlink)

        return hyperlink

    @staticmethod
    def define_color_of_sentiment_cell(value):
        match value.text.strip():
            case 'Нейтральная' | 'Бейтарап' | 'Neutral':
                shading_elm = parse_xml(r'<w:shd {} w:fill="#FFFF00"/>'.format(nsdecls('w')))
                value._tc.get_or_add_tcPr().append(shading_elm)
            case 'Негативная' | 'Негатив' | 'Negative':
                shading_elm = parse_xml(r'<w:shd {} w:fill="#FF0000"/>'.format(nsdecls('w')))
                value._tc.get_or_add_tcPr().append(shading_elm)
            case 'Позитивная' | 'Позитив' | 'Positive':
                shading_elm = parse_xml(r'<w:shd {} w:fill="#008000"/>'.format(nsdecls('w')))
                value._tc.get_or_add_tcPr().append(shading_elm)

    @staticmethod
    def highlight_tag(run, paragraph, tags, column_name, tags_highlight_settings):
        """ Highlight для тегов. """

        runs_to_remove = []

        column_name = column_name.strip(':')

        if any(element in run.text.lower() for element in tags) and column_name in (
                'Краткое содержание', 'Пост', 'Қысқаша мазмұны', 'Summary', 'Post'
        ):

            pattern = r'({})'.format('|'.join(map(re.escape, tags)))
            split_parts = re.split(pattern, run.text.lower())

            back_color = tags_highlight_settings.get('back_color')

            runs_to_remove.append(run)
            for i, part in enumerate(split_parts):
                new_run = paragraph.add_run(part + ' ')

                for tag in tags:
                    if tag.lower() in part.lower():

                        TableStylesGenerator.apply_run_styles(new_run, tags_highlight_settings)

                        if back_color:
                            tag = new_run._r
                            shd = OxmlElement('w:shd')
                            shd.set(qn('w:val'), 'clear')
                            shd.set(qn('w:color'), 'auto')
                            shd.set(qn('w:fill'), back_color)
                            tag.rPr.append(shd)

                new_run.font.size = Pt(10)
                new_run.font.name = 'Arial'

        for old_run in runs_to_remove:
            paragraph._p.remove(old_run._r)

    @staticmethod
    def apply_run_styles(run, setting):
        bold = setting.get('bold')
        italic = setting.get('italic')
        underline = setting.get('underline')
        font_color = setting.get('color')

        if bold:
            run.font.bold = bold
        if italic:
            run.font.italic = italic
        if underline:
            run.font.underline = underline
        if font_color:
            red = int(font_color[1:3], 16)
            green = int(font_color[3:5], 16)
            blue = int(font_color[5:7], 16)
            run.font.color.rgb = RGBColor(red, green, blue)

        run.font.size = Pt(10)
        run.font.name = 'Arial'

    @tricky_loggy
    def delete_first_row(self, table):
        if self.pointer > 0 and len(table.rows) > 1:
            tr = table.rows[0]._tr
            table._tbl.remove(tr)

