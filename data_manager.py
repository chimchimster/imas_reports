import os
import re
import threading

from docx.opc.oxml import nsmap
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx.shared import RGBColor, Pt, Cm
from datetime import datetime
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from operator import itemgetter
import docx
from threading import Thread
from docxcompose.composer import Composer
from docxtpl import DocxTemplate


class MergeReport:
    path_to_folder = os.getcwd() + '/temp/'
    path_to_templates = os.getcwd() + '/templates/'
    path_to_result = os.getcwd() + '/result/'

    def merge(self):

        master = docx.Document(self.path_to_templates + 'out.docx')
        composer = Composer(master)

        file_order = [file for file in os.listdir(self.path_to_folder)]
        file_order.sort()

        page_break_added = False

        for idx, file in enumerate(file_order):

            file_path = os.path.join(self.path_to_folder, file)

            if os.path.isfile(file_path) and file.endswith('.docx'):
                doc = docx.Document(file_path)

                if file_path.endswith('table.docx') and not page_break_added:
                    run = master.add_paragraph().add_run()
                    run.add_break(docx.enum.text.WD_BREAK.PAGE)
                    page_break_added = True

                composer.append(doc)

        output_file = os.path.join(self.path_to_result, 'merged_output.docx')
        composer.save(output_file)


class ThreadDataGenerator(Thread):

    templates = {
        'table': 'templates/template_parts/table.docx',
        'table_of_contents': 'templates/template_parts/table_of_contents.docx',
        'tags': 'templates/template_parts/tags.docx',
    }

    def __init__(self, thread_obj):
        super().__init__()
        self.thread_obj = thread_obj

    def run(self) -> None:

        self.thread_obj.generate_data()

        data = self.thread_obj.data_collection

        if self.thread_obj.flag == 'table':

            template = DocxTemplate(self.templates['table'])

            position = self.thread_obj._rest_data.get('position')

            output_path = f'temp/output-{position}-table.docx'

            table_name = self.thread_obj._rest_data.get('id')
            table = self.thread_obj._rest_data.get('table')

            try:
                template.render({'columns': data[0].keys(), 'data': data, 'is_table': table, 'table_name': table_name}, autoescape=True)
            except IndexError:
                raise IndexError('Невозможно сформировавть таблицу по причине нехватки данных!')

            settings = {k: v for (k, v) in self.thread_obj._rest_data.items()}

            tags = self.thread_obj._data.get('query_ar')

            tags_highlight_settings = self.thread_obj._rest_data.get('tag_highlight')

            if self.thread_obj._rest_data.get('table'):
                styles = TableStylesGenerator(template, tags,  settings, tags_highlight_settings)
                styles.apply_table_styles()
            else:
                styles = SchedulerStylesGenerator(template, tags, settings, tags_highlight_settings)
                styles.apply_scheduler_styles()

            template.save(output_path)

        elif self.thread_obj.flag == 'content':

            template = DocxTemplate(self.templates['table_of_contents'])

            position = self.thread_obj._rest_data.get('position')

            output_path = f'temp/output-{position}-content.docx'

            has_soc = data.get('soc', {})
            has_smi = data.get('smi', {})

            try:
                template.render({'table_of_contents_soc': has_soc, 'table_of_contents_smi': has_smi}, autoescape=True)
            except IndexError:
                raise IndexError('Невозможно сформировать оглавление по причине нехватки данных!')

            template.save(output_path)

        elif self.thread_obj.flag == 'tags':

            template = DocxTemplate(self.templates['tags'])

            position = self.thread_obj._rest_data.get('position')

            output_path = f'temp/output-{position}-atags.docx'

            try:
                template.render({'tags': data}, autoescape=True)
            except IndexError:
                raise IndexError('Невозможно сформировать теги по причине нехватки данных!')
            template.save(output_path)


class DataManager:
    def __init__(self, rest_data, result_data):
        self._rest_data = rest_data
        self._result_data = result_data
        self.threads_objs = []

    def distribute_content(self):
        for data in self._rest_data:
            match data.get('id'):
                case 'tags':
                    tags_obj = TagsGenerator(self._result_data, data)
                    self.threads_objs.append(tags_obj)
                case 'contents':
                    contents_obj = ContentGenerator(self._result_data, data)
                    self.threads_objs.append(contents_obj)
                case 'smi':
                    smi_table_obj = TableContentGenerator(self._result_data, data, 'smi')
                    self.threads_objs.append(smi_table_obj)
                case 'soc':
                    soc_table_obj = TableContentGenerator(self._result_data, data, 'soc')
                    self.threads_objs.append(soc_table_obj)

    def apply_threads(self):

        threads = []

        for thread_obj in self.threads_objs:
            thread = ThreadDataGenerator(thread_obj)
            threads.append(thread)
            thread.start()

        for thr in threads:
            thr.join()


class ContentGenerator:

    flag = 'content'

    def __init__(self, data, rest_data):
        self._data = data
        self._rest_data = rest_data
        self.data_collection = {'soc': [], 'smi': []}

    def generate_data(self):

        cut = 150

        def collect_titles_or_texts(news, counter, key):

            if counter > 50:
                counter = 50

            for idx in range(counter):
                try:
                    text_obj = news[idx][key][:cut].strip()

                    self.data_collection.append(text_obj + ' ...' if len(text_obj) == cut else text_obj)
                except IndexError:
                    pass

        def check_length_of_title_or_text(post):

            post = post.strip()

            if len(post) > 150:
                return post[:cut] + ' ...'

            return post

        if self._rest_data.get('id') == 'contents':
            count_soc = self._rest_data.get('soc')
            count_smi = self._rest_data.get('smi')

            soc_posts = self._data.get('f_news2')
            smi_posts = self._data.get('f_news')

            if count_smi > 0:
                collect_titles_or_texts(smi_posts, count_smi, 'title')
            else:

                for post in smi_posts:
                    self.data_collection['smi'].append(check_length_of_title_or_text(post.get('title')))

            if count_soc > 0:
                collect_titles_or_texts(soc_posts, count_soc, 'full_text')
            else:

                for post in soc_posts:
                    self.data_collection['soc'].append(check_length_of_title_or_text(post.get('full_text')))

        for key, value in self.data_collection.items():
            self.data_collection[key] = {k: v for (k, v) in enumerate(value, start=1)}


class ContentAnchors:
    def __init__(self, document):
        self.document = document

    def apply_anchors(self):
        for paragraph in self.document.paragraphs:
            for run in paragraph.runs:
                print(run.text)

        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    print(cell.text)

doc_path = '../result/merged_output.docx'
doc = docx.Document(doc_path)

c = ContentAnchors(doc)
c.apply_anchors()


class TagsGenerator:

    flag = 'tags'

    def __init__(self, data, rest_data):
        self._data = data
        self._rest_data = rest_data
        self.data_collection = []

    def generate_data(self):

        self.data_collection.append(self._data.get('analyzer_tags_changed'))


class TableContentGenerator:

    flag = 'table'

    translator_smi = {
        'title': 'Заголовок',
        'not_date': 'Дата',
        'full_text': 'Краткое содержание',
        'RESOURCE_NAME': 'Наименование СМИ',
        'news_link': 'URL',
        'sentiment': 'Тональность',
        'name_cat': 'Категория',
        'number': '№',
    }

    translator_soc = {
        'date': 'Дата',
        'full_text': 'Пост',
        'resource_name': 'Сообщество',
        'news_link': 'URL',
        'sentiment': 'Тональность',
        'type': 'Соцсеть',
        'number': '№',
    }

    def __init__(self, data, rest_data, _type):
        self._data = data
        self._rest_data = rest_data
        self._type = _type
        self.data_collection = []

    def generate_data(self):

        soc_types = {
            1: 'Вконтакте',
            2: 'Facebook',
            3: 'Twitter',
            4: 'Instagram',
            5: 'LinkedIn',
            6: 'Youtube',
            7: 'Одноклассники',
            8: 'Мой Мир',
            9: 'Telegram',
            10: 'TikTok',
        }

        def sort_data(table_data):
            order = self._rest_data.get('order')

            if not order:
                return

            date = order.get('date')
            predominantly = order.get('predominantly')
            sentiments = order.get('sentiments')
            categories = order.get('categories')

            def delete_unused_sentiments(_table_data):

                if not _table_data:
                    return []

                if sentiments:
                    for idx, sentiment in enumerate(sentiments):
                        if idx == 0 and sentiment == 0:
                            _table_data = [table for table in _table_data if table['sentiment'] != 1]
                        elif idx == 1 and sentiment == 0:
                            _table_data = [table for table in _table_data if table['sentiment'] != -1]
                        elif idx == 2 and sentiment == 0:
                            _table_data = [table for table in _table_data if table['sentiment'] != 0]

                    return _table_data

                return _table_data

            def delete_unused_categories(_table_data):

                if not _table_data:
                    return []

                if categories:
                    try:
                        return [table for table in _table_data if table['name_cat'] in categories]
                    except:
                        return [table for table in _table_data if soc_types[table['type']] in categories]

                return _table_data

            def sort_by_sentiment_category_date(_table_data):

                if not _table_data:
                    return []

                sentiment_index = {1: 0, 0: 1, -1: 2}

                category_index = {category: i for i, category in enumerate(categories)}

                if _table_data[0].get('name_cat'):
                    return sorted(_table_data, key=lambda x: (
                        sentiment_index[x['sentiment']], category_index.get(x['name_cat'], len(categories)), x['nd_date']),
                                  reverse=date == 0)
                elif _table_data[0].get('type'):
                    return sorted(_table_data, key=lambda x: (
                        sentiment_index[x['sentiment']], category_index.get(x['type'], len(categories)), x['date']),
                                  reverse=date == 0)

                return _table_data

            def sort_by_category_sentiment_date(_table_data):

                if not _table_data:
                    return []

                sentiment_index = {1: 0, 0: 1, -1: 2}

                category_index = {category: i for i, category in enumerate(categories)}

                if _table_data[0].get('name_cat'):
                    return sorted(_table_data, key=lambda x: (
                        category_index.get(x['name_cat'], len(categories)), sentiment_index[x['sentiment']], x['nd_date']),
                                  reverse=date == 0)
                elif _table_data[0].get('type'):
                    return sorted(_table_data, key=lambda x: (
                        category_index.get(x['type'], len(categories)), sentiment_index[x['sentiment']], x['date']),
                                  reverse=date == 0)

                return _table_data

            def sort_by_sentiment_date(_table_data):

                if not _table_data:
                    return []

                if _table_data[0].get('nd_date'):
                    return sorted(_table_data, key=lambda x: (x['sentiment'], x['nd_date']), reverse=date == 0)
                elif _table_data[0].get('date'):
                    return sorted(_table_data, key=lambda x: (x['sentiment'], x['date']), reverse=date == 0)

                return _table_data

            def sort_by_category_date(_table_data):

                if not _table_data:
                    return []

                if _table_data[0].get('nd_date'):
                    return sorted(_table_data, key=lambda x: (x['name_cat'], x['nd_date']), reverse=date == 0)
                elif _table_data[0].get('date'):
                    return sorted(_table_data, key=lambda x: (x['type'], x['date']), reverse=date == 0)

                return _table_data

            def sort_by_date(_table_data):

                if not _table_data:
                    return []

                if _table_data[0].get('nd_date'):
                    return sorted(_table_data, key=itemgetter('nd_date'), reverse=date == 0)
                elif _table_data[0].get('date'):
                    return sorted(_table_data, key=itemgetter('date'), reverse=date == 0)

                return _table_data

            table_data = delete_unused_sentiments(table_data)

            table_data = delete_unused_categories(table_data)

            if sentiments and categories:
                if predominantly == 0:
                    sorted_table_data = sort_by_sentiment_category_date(table_data)
                else:
                    sorted_table_data = sort_by_category_sentiment_date(table_data)
            elif sentiments != 0:
                sorted_table_data = sort_by_sentiment_date(table_data)
            elif categories != 0:
                sorted_table_data = sort_by_category_date(table_data)
            else:
                sorted_table_data = sort_by_date(table_data)

            return sorted_table_data

        if self._type == 'smi':
            f_news = self._data.get('f_news')
            f_news = sort_data(f_news)
            if f_news:
                self.__apply_translator(self.translator_smi, f_news)
        else:
            f_news2 = self._data.get('f_news2')
            f_news2 = sort_data(f_news2)
            if f_news2:
                self.__apply_translator(self.translator_soc, f_news2)

    def __apply_translator(self, translator, news):

        translator_for_rest_soc = {
            'number': 'number',
            'content': 'full_text',
            'date': 'date',
            'resource': 'resource_name',
            'news_link': 'news_link',
            'sentiment': 'sentiment',
            'category': 'type',
        }

        translator_for_rest_smi = {
            'number': 'number',
            'title': 'title',
            'content': 'full_text',
            'date': 'not_date',
            'resource': 'RESOURCE_NAME',
            'news_link': 'news_link',
            'sentiment': 'sentiment',
            'category': 'name_cat',
        }

        to_sort = {}
        to_delete = []

        def delete_unused_columns(_table, translator_type):
            for tbl in _table['columns']:
                column_name = tbl.get('id') if tbl.get('position') == 0 else None
                if column_name:
                    to_delete.append(translator_type[column_name])

        def sort_columns(_table, translator_type):
            for tbl in _table['columns']:
                to_sort[tbl.get('id')] = tbl['position']
            return {translator[translator_type[k]]: v for (k, v) in to_sort.items()}

        def update_collection():

            def choose_tag(tags, value_string):

                for tag in tags:
                    if tag.lower() in value_string.lower():
                        return tag.lower()
                return ''

            text_length = self._rest_data.get('text_length')
            tags = self._data.get('query_ar')

            for i in range(len(news)):
                news[i] = {**{'number': i + 1}, **news[i]}

                if news[i].get('date'):
                    news[i]['date'] = datetime.fromtimestamp(news[i]['date']).strftime('%d-%m-%Y')

                if news[i].get('type'):
                    self.__match_social_medias(news[i])

                result = {}

                for key, value in news[i].items():

                    value = value.strip() if isinstance(value, str) else value

                    if key in translator:
                        if translator[key] in ('Пост', 'Краткое содержание'):
                            tag = choose_tag(tags, value)
                            temp_val = value.lower()

                            if len(value) <= text_length:
                                result[translator[key]] = value
                                continue

                            if tag == '':
                                result[translator[key]] = value[:text_length] + ' ...' if text_length < len(value) else value[:text_length]
                                continue

                            tag_start = temp_val.find(tag)
                            if tag_start != -1:
                                tag_end = tag_start + len(tag)
                                left = max(0, tag_start - (text_length - len(tag)) // 2)
                                right = min(len(value), tag_end + (text_length - len(tag)) // 2)
                                result[translator[key]] = '...' + value[left:right] + '...'
                        else:
                            result[translator[key]] = value

                sorted_result = {k: v for (k, v) in sorted(result.items(), key=lambda x: to_sort[x[0]])}
                self.data_collection.append(sorted_result)

        if self._rest_data.get('id') == 'soc':
            delete_unused_columns(self._rest_data, translator_for_rest_soc)
        elif self._rest_data.get('id') == 'smi':
            delete_unused_columns(self._rest_data, translator_for_rest_smi)

        news = [{k: v for (k, v) in n.items() if k not in to_delete} for n in news]

        if self._rest_data.get('id') == 'soc':
            to_sort = sort_columns(self._rest_data, translator_for_rest_soc)
        elif self._rest_data.get('id') == 'smi':
            to_sort = sort_columns(self._rest_data, translator_for_rest_smi)

        update_collection()

    def __match_social_medias(self, data):
        match data.get('type'):
            case 1:
                data['type'] = 'Вконтакте'
            case 2:
                data['type'] = 'Facebook'
            case 3:
                data['type'] = 'Twitter'
            case 4:
                data['type'] = 'Instagram'
            case 5:
                data['type'] = 'LinkedIn'
            case 6:
                data['type'] = 'Youtube'
            case 7:
                data['type'] = 'Одноклассники'
            case 8:
                data['type'] = 'Мой Мир'
            case 9:
                data['type'] = 'Telegram'
            case 10:
                data['type'] = 'TikTok'


class TableStylesGenerator:

    translator_smi = {
        'title': 'Заголовок',
        'date': 'Дата',
        'content': 'Краткое содержание',
        'resource': 'Наименование СМИ',
        'news_link': 'URL',
        'sentiment': 'Тональность',
        'category': 'Категория',
        'number': '№',
    }

    translator_soc = {
        'date': 'Дата',
        'content': 'Пост',
        'resource': 'Сообщество',
        'news_link': 'URL',
        'sentiment': 'Тональность',
        'category': 'Соцсеть',
        'number': '№',
    }

    def __init__(self, template, tags, settings=None, tags_highlight_settings=None):
        self._template = template
        self._tags = tags
        self._settings = settings
        self._tags_highlight_settings = tags_highlight_settings

    def apply_table_styles(self):

        table = self._template.tables[0]
        table.style = 'Table Grid'

        if not self._settings:
            return

        if self._settings.get('id') == 'soc':
            self.choose_particular_table_styles(self.translator_soc, table, 'soc')
        else:
            self.choose_particular_table_styles(self.translator_smi, table, 'smi')

    def choose_particular_table_styles(self, translator_obj, table_obj, _type):
        def set_cell_width():
            match cell.text:
                case '№':
                    table_obj.columns[idx].width = Cm(1)
                case 'Заголовок':
                    table_obj.columns[idx].width = Cm(8)
                case "Пост" | 'Краткое содержание':
                    table_obj.columns[idx].width = Cm(15)
                case 'Дата':
                    table_obj.columns[idx].width = Cm(5)
                case 'Соцсеть' | 'Категория':
                    table_obj.columns[idx].width = Cm(5)
                case 'URL':
                    table_obj.columns[idx].width = Cm(5)
                case 'Сообщество' | 'Наименование СМИ':
                    table_obj.columns[idx].width = Cm(8)
                case 'Тональность':
                    table_obj.columns[idx].width = Cm(6)

        for column in self._settings['columns']:
            if column.get('id') in translator_obj:
                column_name_en = column.get('id')
                column_name = translator_obj[column_name_en]

                for idx, cell in enumerate(table_obj.row_cells(0)):
                    set_cell_width()

                    if cell.text == column_name:
                        for row in table_obj.rows[1:]:
                            cell = row.cells[idx]

                            self.define_color_of_sentiment_cell(cell)

                            for paragraph in cell.paragraphs:
                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                for run in paragraph.runs:

                                    self.highlight_tag(run, paragraph, self._tags, column_name, self._tags_highlight_settings)

                                    if re.match(r'https?://(?:[-\w.]|(?:%[\da-fA-F]{2}))+[/\w .?=-]*', cell.text) and column_name == 'URL':
                                        hyperlink = self.add_hyperlink(paragraph, cell.text.strip(), 'Ссылка', '#0000FF', '#000080')

                                        for old_run in paragraph.runs:
                                            if old_run != hyperlink:
                                                paragraph._p.remove(old_run._r)

                                    self.apply_run_styles(run, column)

    @staticmethod
    def add_hyperlink(paragraph, url, text, color, underline):

        part = paragraph.part
        r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

        new_run = docx.oxml.shared.OxmlElement('w:r')

        rPr = docx.oxml.shared.OxmlElement('w:rPr')

        if not color is None:
            c = docx.oxml.shared.OxmlElement('w:color')
            c.set(docx.oxml.shared.qn('w:val'), color)
            rPr.append(c)

        if not underline:
            u = docx.oxml.shared.OxmlElement('w:u')
            u.set(docx.oxml.shared.qn('w:val'), 'none')
            rPr.append(u)

        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)

        paragraph.add_run(' ')
        paragraph._p.append(hyperlink)

        return hyperlink

    @staticmethod
    def define_color_of_sentiment_cell(value):
        match value.text.strip():
            case 'Нейтральная':
                shading_elm = parse_xml(r'<w:shd {} w:fill="#FFFF00"/>'.format(nsdecls('w')))
                value._tc.get_or_add_tcPr().append(shading_elm)
            case 'Негативная':
                shading_elm = parse_xml(r'<w:shd {} w:fill="#FF0000"/>'.format(nsdecls('w')))
                value._tc.get_or_add_tcPr().append(shading_elm)
            case 'Позитивная':
                shading_elm = parse_xml(r'<w:shd {} w:fill="#008000"/>'.format(nsdecls('w')))
                value._tc.get_or_add_tcPr().append(shading_elm)

    @staticmethod
    def highlight_tag(run, paragraph, tags, column_name, tags_highlight_settings):
        """ Highlight для тегов. """

        runs_to_remove = []

        column_name = column_name.strip(':')

        if any(element in run.text.lower() for element in tags) and column_name in ('Краткое содержание', 'Пост'):

            pattern = r"\b" + r"\b|\b".join(map(re.escape, tags)) + r"\b"
            split_parts = re.split(f"({pattern})", run.text.lower())

            back_color = tags_highlight_settings.get('back_color')

            runs_to_remove.append(run)
            for i, part in enumerate(split_parts):
                new_run = paragraph.add_run(part + ' ')

                for tag in tags:
                    if tag.lower() in part.lower():

                        TableStylesGenerator.apply_run_styles(new_run, tags_highlight_settings)

                        if back_color:
                            tag = new_run._r
                            shd = OxmlElement('w:shd')
                            shd.set(qn('w:val'), 'clear')
                            shd.set(qn('w:color'), 'auto')
                            shd.set(qn('w:fill'), back_color)
                            tag.rPr.append(shd)

                new_run.font.size = Pt(10)
                new_run.font.name = 'Arial'

        for old_run in runs_to_remove:
            paragraph._p.remove(old_run._r)

    @staticmethod
    def apply_run_styles(run, setting):
        bold = setting.get('bold')
        italic = setting.get('italic')
        underline = setting.get('underline')
        font_color = setting.get('color')

        if bold:
            run.font.bold = bold
        if italic:
            run.font.italic = italic
        if underline:
            run.font.underline = underline
        if font_color:
            red = int(font_color[1:3], 16)
            green = int(font_color[3:5], 16)
            blue = int(font_color[5:7], 16)
            run.font.color.rgb = RGBColor(red, green, blue)

        run.font.size = Pt(10)
        run.font.name = 'Arial'


class SchedulerStylesGenerator(TableStylesGenerator):

    def __init__(self, template, tags, settings=None, tags_highlight_settings=None):
        super().__init__(template, tags, settings, tags_highlight_settings)

    def apply_scheduler_styles(self):

        def manage_styles(paragraph, curr_run, prev_run, rows):

            prev_run_text = prev_run.text.rstrip(':')

            def get_setting():
                for row in rows:
                    id = row.get('id')
                    if self.translator_smi.get(id) == prev_run_text or self.translator_soc.get(id) == prev_run_text:
                        return row
                return None

            setting = get_setting()

            if setting:
                self.apply_run_styles(curr_run, setting)

            if prev_run_text in ('Заголовок', 'Краткое содержание', 'Пост'):
                paragraph._p.remove(prev_run._r)

            if prev_run_text == '№':
                paragraph._p.remove(curr_run._r)
                paragraph._p.remove(prev_run._r)

        if not self._settings:
            return

        scheduler = self._template.paragraphs

        prev_run = None
        rows = self._settings['list_rows']
        for paragraph in scheduler:
            for idx, run in enumerate(paragraph.runs, start=1):
                curr_run = run

                if prev_run:
                    self.highlight_tag(curr_run, paragraph, self._tags, prev_run.text, self._tags_highlight_settings)

                if re.match(r'https?://(?:[-\w.]|(?:%[\da-fA-F]{2}))+[/\w .?=-]*', curr_run.text.strip()):
                    if prev_run.text.strip(':') == 'URL':
                        self.add_hyperlink(paragraph, curr_run.text.strip(), 'Ссылка', '#0000FF', '#000080')

                        paragraph._p.remove(curr_run._r)

                if idx % 2 == 0:
                    manage_styles(paragraph, curr_run, prev_run, rows)
                else:
                    prev_run = run
