import re
import json
from collections import defaultdict
from copy import deepcopy
from itertools import groupby
from operator import itemgetter

import docx
import requests
from docx.opc.oxml import nsmap
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx.shared import Mm, RGBColor, Pt, Cm
from docxtpl import DocxTemplate, InlineImage
import jinja2
from datetime import datetime
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_COLOR_INDEX


class HTMLWordCreator:
    api_url = 'https://rest.imas.kz/'
    relative_api = 'export-apis?'
    relative_api_folders = 'export-api-folders?'

    def __init__(self, rest_data):
        self._rest_data = rest_data

    def render_report(self):

        # for i in self._rest_data:
        #     print(i, end='\n')

        response_string = self.generate_string()

        an_id = self._rest_data[-1].get('an_id')

        query_url = self.__define_query_url(an_id)

        response = requests.get(query_url + response_string)

        if response.status_code == 200:
            result = response.json()

            self.generate_word_document(result)

    def generate_word_document(self, result):
        template_path = 'templates/template_parts/table.docx'
        output_path = 'output.docx'

        template = DocxTemplate(template_path)

        data = TableContentGenerator(result, self._rest_data, type='smi')
        data.generate_data()

        data = data.data_collection
        template.render({'columns': data[0].keys(), 'data': data}, autoescape=True)

        settings = [_data for _data in self._rest_data if _data.get('id') in ('smi', 'soc')]

        tags = result.get('query_ar')
        tags_highlight_settings = self._rest_data[0].get('tag_highlight')

        styles = TableStylesGenerator(template, tags,  settings, tags_highlight_settings)
        styles.apply_table_styles()

        template.save(output_path)


    def generate_string(self) -> str:
        """ Формирование строки из словаря пришедшего в POST-запросе. """

        request_string = ''

        self._rest_data[-1]['format'] = 'pdf'
        self._rest_data[-1]['location'] = 2
        self._rest_data[-1]['full_text'] = 1

        request_string += '&'.join(['='.join([key, str(val)]) for (key, val) in self._rest_data[-1].items()])

        return request_string

    @classmethod
    def __define_query_url(cls, value: str) -> str:
        """ Определяем строку для запроса. """

        try:
            int(value)
            return cls.api_url + cls.relative_api
        except:
            return cls.api_url + cls.relative_api_folders

    @staticmethod
    def parse_string(query_string: str) -> dict:
        """ Формирование словаря из данных пришедших в виде строки с REST. """

        return dict(re.findall(r'([^&=]+)=([^&]*)', query_string))


class TableContentGenerator:
    translator_smi = {
        'title': 'Заголовок',
        'not_date': 'Дата',
        'content': 'Краткое содержание',
        'RESOURCE_NAME': 'Наименование СМИ',
        'news_link': 'URL',
        'sentiment': 'Тональность',
        'name_cat': 'Категория',
        'number': '№',
    }

    translator_soc = {
        'date': 'Дата',
        'text': 'Пост',
        'RESOURCE_NAME': 'Сообщество',
        'news_link': 'URL',
        'sentiment': 'Тональность',
        'type': 'Соцсеть',
        'number': '№',
    }

    def __init__(self, data, rest_data, type):
        self._data = data
        self._type = type
        self._rest_data = rest_data
        self.data_collection = []

    def generate_data(self):

        def sort_data(table_data):
            order = self._rest_data[0].get('order')

            if not order:
                return

            date = order.get('date')
            predominantly = order.get('predominantly')
            sentiments = order.get('sentiments')
            categories = [category['name_cat'] for category in order.get('categories', []) if isinstance(category, dict)]

            def sort_by_sentiment_category_date(table_data):
                sentiment_index = {1: 0, 0: 1, -1: 2}

                category_index = {category: i for i, category in enumerate(categories)}

                return sorted(table_data, key=lambda x: (sentiment_index[x['sentiment']], category_index.get(x['name_cat'], len(categories)), x['nd_date']), reverse=date == 0)

            def sort_by_category_sentiment_date(table_data):
                sentiment_index = {1: 0, 0: 1, -1: 2}

                category_index = {category: i for i, category in enumerate(categories)}

                return sorted(table_data, key=lambda x: (category_index.get(x['name_cat'], len(categories)), sentiment_index[x['sentiment']], x['nd_date']), reverse=date == 0)

            def sort_by_sentiment_date(table_data):

                return sorted(table_data, key=lambda x: (x['sentiment'], x['nd_date']), reverse=date == 0)

            def sort_by_category_date(table_data):

                return sorted(table_data, key=lambda x: (x['name_cat'], x['nd_date']), reverse=date == 0)

            def sort_by_date(table_data):

                return sorted(table_data, key=itemgetter('nd_date'), reverse=date == 0)

            if sentiments and categories:
                if predominantly == 0:
                    sorted_table_data = sort_by_sentiment_category_date(table_data)
                else:
                    sorted_table_data = sort_by_category_sentiment_date(table_data)
            elif sentiments:
                sorted_table_data = sort_by_sentiment_date(table_data)
            elif categories:
                sorted_table_data = sort_by_category_date(table_data)
            else:
                sorted_table_data = sort_by_date(table_data)

            return sorted_table_data

        if self._type == 'smi':
            f_news = self._data.get('f_news')
            f_news = sort_data(f_news)
            if f_news:
                self.__apply_translator(self.translator_smi, f_news)
        else:
            f_news2 = self._data.get('f_news2')
            if f_news2:
                self.__apply_translator(self.translator_soc, f_news2)

    def __apply_translator(self, translator, news):

        translator_for_rest_soc = {
            'number': 'number',
            'content': 'text',
            'date': 'date',
            'resource': 'resource_name',
            'news_link': 'news_link',
            'sentiment': 'sentiment',
            'category': 'type',
        }

        translator_for_rest_smi = {
            'number': 'number',
            'title': 'title',
            'content': 'content',
            'date': 'not_date',
            'resource': 'RESOURCE_NAME',
            'news_link': 'news_link',
            'sentiment': 'sentiment',
            'category': 'name_cat',
        }

        to_sort = {}
        to_delete = []

        def delete_unused_columns(_table, translator_type):
            for tbl in _table['columns']:
                column_name = tbl.get('id') if tbl.get('position') == 0 else None
                if column_name:
                    to_delete.append(translator_type[column_name])

        def sort_columns(_table, translator_type):
            for tbl in _table['columns']:
                to_sort[tbl.get('id')] = tbl['position']

            return {translator[translator_type[k]]: v for (k, v) in to_sort.items()}

        def update_collection():
            for i in range(len(news)):
                news[i] = {**{'number': i + 1}, **news[i]}

                if news[i].get('date'):
                    news[i]['date'] = datetime.fromtimestamp(news[i]['date']).strftime('%d-%m-%Y')

                if news[i].get('type'):
                    self.__match_social_medias(news[i])

                result = {translator[k]: v for (k, v) in news[i].items() if k in translator}
                sorted_result = {k: v for (k, v) in sorted(result.items(), key=lambda x: to_sort[x[0]])}
                self.data_collection.append(sorted_result)

        # Удаление ненужных столбцов
        for table in self._rest_data:
            if table.get('id') == 'soc':
                delete_unused_columns(table, translator_for_rest_soc)
            elif table.get('id') == 'smi':
                delete_unused_columns(table, translator_for_rest_smi)

        news = [{k: v for (k, v) in n.items() if k not in to_delete} for n in news]

        # Сортировка столбцов
        for table in self._rest_data:
            if table.get('id') == 'soc':
                to_sort = sort_columns(table, translator_for_rest_soc)
            elif table.get('id') == 'smi':
                to_sort = sort_columns(table, translator_for_rest_smi)

        update_collection()

    def __match_social_medias(self, data):
        match data.get('type'):
            case 1:
                data['type'] = 'Вконтакте'
            case 2:
                data['type'] = 'Facebook'
            case 3:
                data['type'] = 'Twitter'
            case 4:
                data['type'] = 'Instagram'
            case 5:
                data['type'] = 'LinkedIn'
            case 6:
                data['type'] = 'Youtube'
            case 7:
                data['type'] = 'Одноклассники'
            case 8:
                data['type'] = 'Мой Мир'
            case 9:
                data['type'] = 'Telegram'
            case 10:
                data['type'] = 'TikTok'


class TableStylesGenerator:
    translator_smi = {
        'title': 'Заголовок',
        'date': 'Дата',
        'content': 'Краткое содержание',
        'resource': 'Наименование СМИ',
        'news_link': 'URL',
        'sentiment': 'Тональность',
        'category': 'Категория',
        'number': '№',
    }

    translator_soc = {
        'date': 'Дата',
        'content': 'Пост',
        'resource': 'Сообщество',
        'news_link': 'URL',
        'sentiment': 'Тональность',
        'category': 'Соцсеть',
        'number': '№',
    }

    def __init__(self, template, tags, settings=None, tags_highlight_settings=None):
        self._template = template
        self._tags = tags
        self._settings = settings
        self._tags_highlight_settings = tags_highlight_settings

    def apply_table_styles(self):

        table = self._template.tables[0]
        table.style = 'Table Grid'

        if not self._settings:
            return

        if self._settings[0].get('id') == 'soc':
            self.choose_particular_table_styles(self.translator_soc, table, 'soc')
        else:
            self.choose_particular_table_styles(self.translator_smi, table, 'smi')

    def choose_particular_table_styles(self, translator_obj, table_obj, _type):
        def set_cell_width():
            match cell.text:
                case '№':
                    table_obj.columns[idx].width = Cm(1)
                case 'Заголовок':
                    table_obj.columns[idx].width = Cm(8)
                case "Пост" | 'Краткое содержание':
                    table_obj.columns[idx].width = Cm(15)
                case 'Дата':
                    table_obj.columns[idx].width = Cm(5)
                case 'Соцсеть' | 'Категория':
                    table_obj.columns[idx].width = Cm(5)
                case 'URL':
                    table_obj.columns[idx].width = Cm(5)
                case 'Сообщество' | 'Наименование СМИ':
                    table_obj.columns[idx].width = Cm(8)
                case 'Тональность':
                    table_obj.columns[idx].width = Cm(6)

        for setting in self._settings:
            for column in setting['columns']:
                if column.get('id') in translator_obj:
                    column_name_en = column.get('id')
                    column_name = translator_obj[column_name_en]

                    for idx, cell in enumerate(table_obj.row_cells(0)):
                        set_cell_width()

                        if cell.text == column_name:
                            for row in table_obj.rows[1:]:
                                cell = row.cells[idx]

                                self.define_color_of_sentiment_cell(cell)

                                for paragraph in cell.paragraphs:
                                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                    for run in paragraph.runs:

                                        self.highlight_tag(run, paragraph, self._tags, column_name, self._tags_highlight_settings)

                                        if re.match(r'https?://(?:[-\w.]|(?:%[\da-fA-F]{2}))+[/\w .?=-]*', cell.text):
                                            hyperlink = self.add_hyperlink(paragraph, cell.text.strip(), 'Ссылка', '#0000FF', '#000080')

                                            for old_run in paragraph.runs:
                                                if old_run != hyperlink:
                                                    paragraph._p.remove(old_run._r)

                                        bold = column.get('bold')
                                        italic = column.get('italic')
                                        underline = column.get('underline')
                                        color = column.get('color')

                                        if bold:
                                            run.font.bold = bold

                                        if italic:
                                            run.font.italic = italic

                                        if underline:
                                            run.font.underline = underline

                                        if color:
                                            red = int(color[1:3], 16)
                                            green = int(color[3:5], 16)
                                            blue = int(color[5:7], 16)
                                            run.font.color.rgb = RGBColor(red, green, blue)

                                        run.font.name = 'Arial'
                                        run.font.size = Pt(10)

    @staticmethod
    def add_hyperlink(paragraph, url, text, color, underline):

        part = paragraph.part
        r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

        new_run = docx.oxml.shared.OxmlElement('w:r')

        rPr = docx.oxml.shared.OxmlElement('w:rPr')

        if not color is None:
            c = docx.oxml.shared.OxmlElement('w:color')
            c.set(docx.oxml.shared.qn('w:val'), color)
            rPr.append(c)

        if not underline:
            u = docx.oxml.shared.OxmlElement('w:u')
            u.set(docx.oxml.shared.qn('w:val'), 'none')
            rPr.append(u)

        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)

        paragraph._p.append(hyperlink)

        return hyperlink

    @staticmethod
    def define_color_of_sentiment_cell(value):
        match value.text.strip():
            case 'Нейтральная':
                shading_elm = parse_xml(r'<w:shd {} w:fill="#FFFF00"/>'.format(nsdecls('w')))
                value._tc.get_or_add_tcPr().append(shading_elm)
            case 'Негативная':
                shading_elm = parse_xml(r'<w:shd {} w:fill="#FF0000"/>'.format(nsdecls('w')))
                value._tc.get_or_add_tcPr().append(shading_elm)
            case 'Позитивная':
                shading_elm = parse_xml(r'<w:shd {} w:fill="#008000"/>'.format(nsdecls('w')))
                value._tc.get_or_add_tcPr().append(shading_elm)

    @staticmethod
    def highlight_tag(run, paragraph, tags, column_name, tags_highlight_settings):
        """ Highlight для тегов. """

        runs_to_remove = []

        if any(element in run.text.lower() for element in tags) and column_name == 'Краткое содержание':

            pattern = r"\b" + r"\b|\b".join(map(re.escape, tags)) + r"\b"
            split_parts = re.split(f"({pattern})", run.text.lower())

            bold = tags_highlight_settings.get('bold')
            italic = tags_highlight_settings.get('italic')
            underline = tags_highlight_settings.get('underline')
            font_color = tags_highlight_settings.get('color')
            back_color = tags_highlight_settings.get('back_color')

            runs_to_remove.append(run)
            for i, part in enumerate(split_parts):
                new_run = paragraph.add_run(part + ' ')

                for tag in tags:
                    if tag.lower() in part.lower():

                        if bold:
                            new_run.font.bold = bold

                        if italic:
                            new_run.font.italic = italic

                        if underline:
                            new_run.font.underline = underline

                        if font_color:
                            red = int(font_color[1:3], 16)
                            green = int(font_color[3:5], 16)
                            blue = int(font_color[5:7], 16)
                            run.font.color.rgb = RGBColor(red, green, blue)

                        if back_color:
                            tag = new_run._r
                            shd = OxmlElement('w:shd')
                            shd.set(qn('w:val'), 'clear')
                            shd.set(qn('w:color'), 'auto')
                            shd.set(qn('w:fill'), back_color)
                            tag.rPr.append(shd)

                new_run.font.size = Pt(10)
                new_run.font.name = 'Arial'

        for old_run in runs_to_remove:
            paragraph._p.remove(old_run._r)

